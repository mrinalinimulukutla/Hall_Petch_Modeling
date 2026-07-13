#!/usr/bin/env python3
"""
Generate comprehensive analysis report as .docx
Times New Roman 11pt, single space, justified
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import pandas as pd

import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'scripts'))
from _config import (REPO_ROOT, DATA_DIR, RAW_DATA_DIR, RESULTS_DIR, PLOTS_DIR,
                      PAPER_DIR, PAPER_FIG_DIR, REPORT_DIR)
BASE = str(REPO_ROOT)
PLOT_DIR = f'{PLOTS_DIR}'

doc = Document()

# ============================================================
# GLOBAL STYLE SETUP
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.0

# Set default font for East Asian text
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = doc.styles['Normal'].element.get_or_add_rPr().makeelement(qn('w:rFonts'), {})
    rPr.append(rFonts)
rFonts.set(qn('w:ascii'), 'Times New Roman')
rFonts.set(qn('w:hAnsi'), 'Times New Roman')
rFonts.set(qn('w:cs'), 'Times New Roman')

# Margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = rPr.makeelement(qn('w:rFonts'), {})
            rPr.append(rFonts)
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    return h


def add_para(text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    return p


def add_bold_then_text(bold_text, normal_text, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r1 = p.add_run(bold_text)
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(11)
    r1.bold = True
    r2 = p.add_run(normal_text)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(11)
    return p


def add_figure(filename, caption, width=6.0):
    if os.path.exists(os.path.join(PLOT_DIR, filename)):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(os.path.join(PLOT_DIR, filename), width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        cap.paragraph_format.space_after = Pt(12)
        r = cap.add_run(caption)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
        r.italic = True
    else:
        add_para(f'[Figure not found: {filename}]', italic=True)


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.bold = True
    # Data
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
    doc.add_paragraph()  # spacing after table


# ============================================================
# TITLE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(6)
run = title.add_run('Comprehensive Analysis Report')
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(4)
run = subtitle.add_run('Machine Learning-Guided Yield Strength Prediction in FCC\nMulti-Principal Element Alloys: Grain-Size Scaling Laws,\nSolid-Solution Strengthening, and Model Comparison')
run.font.name = 'Times New Roman'
run.font.size = Pt(13)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.paragraph_format.space_after = Pt(18)
run = info.add_run('Al\u2013Co\u2013Cr\u2013Cu\u2013Fe\u2013Mn\u2013Ni\u2013V System | 93 FCC Alloys\nAnalysis Date: March 28\u201329, 2026')
run.font.name = 'Times New Roman'
run.font.size = Pt(11)
run.italic = True

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Dataset Overview',
    '3. Exploratory Data Analysis',
    '4. Solid-Solution Strengthening Models',
    '5. Grain-Size Scaling Laws',
    '6. Bayesian Grain-Size Scaling Analysis',
    '7. Composition-Dependent Hall–Petch Models',
    '8. Exhaustive Machine Learning Model Search',
    '9. SHAP Feature Importance Analysis',
    '10. Symbolic Regression (PySR)',
    '11. SISSO Symbolic Regression and Jiang Model Comparison',
    '12. Information Criteria and Model Selection',
    '13. External Validation Against Independent Literature Data',
    '14. Limitations and Caveats',
    '15. Hardness Analysis',
    '16. Key Findings and Recommendations',
    '17. Related Work',
    'Appendix A: Complete Model Comparison Table',
    'Appendix B: Analysis Plot Catalog',
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
doc.add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
add_heading('1. Executive Summary', level=1)

add_para(
    'This report documents the comprehensive analysis of yield strength prediction in 93 FCC '
    'multi-principal element alloys (MPEAs) spanning the Al\u2013Co\u2013Cr\u2013Cu\u2013Fe\u2013Mn\u2013Ni\u2013V '
    'system. The analysis encompasses exploratory data analysis, physics-based solid-solution strengthening (SSS) '
    'models, nine grain-size scaling laws compared via frequentist (\u0394BIC) methods (eight of '
    'which were also evaluated via Bayesian PSIS-LOO with MCMC), an exhaustive search '
    'over 17 machine learning models with Optuna hyperparameter optimization, SHAP interpretability analysis, '
    'and symbolic regression for equation discovery, plus a comprehensive hardness analysis examining the '
    'Tabor relation and HV Hall\u2013Petch scaling. The work required nine analysis scripts totaling '
    '~6,200 lines of Python code and produced 56 analysis plots, five results tables, and two augmented datasets.'
)

add_para(
    'Key findings: (1) The classical Hall\u2013Petch d\u207b\xb9\u2082 law and three alternatives '
    '(d\u207b\xb2\u2033, ln(d)/d, d\u207b\xb9\u2033) are statistically indistinguishable (\u0394BIC < 2). '
    '(2) All three physics-based SSS models (VLC, Labusch, Toda-Caraballo) produce inaccurate absolute '
    'predictions when implemented with Vegard\'s-law inputs (mean Pred/Exp ratios of 2.2\u201352.8\u00d7); '
    'partial correlations confirm that the Vegard\u2019s-law SSS predictions are approximately linear '
    'functions of composition, carrying negligible nonlinear information beyond elemental fractions '
    '(|r_partial| < 0.1 after controlling for composition). '
    '(3) A composition-dependent \u03c3\u2080 model (\u03c3\u2080 = 230 + \u03a3\u03b1\u1d62\u00b7x\u1d62, '
    'k_HP = 766 MPa\u00b7\u03bcm\xb9\u2082) raises LOO R\u00b2 from 0.41 to 0.65 (Co shear modulus = 75 GPa; '
    'Fe FCC lattice parameter = 3.590 \u00c5). Once \u03c3\u2080(comp) '
    'is properly modeled, no statistically significant composition dependence of k_HP is detected '
    '(all |r| < 0.06, F-test p = 0.999), though the sample size (N = 93) may be insufficient to '
    'detect weak effects. '
    '(4) XGBoost achieves the highest LOO R\u00b2 = 0.729, though its LOBO R\u00b2 = 0.574 reveals '
    'substantial batch-specific overfitting. SISSO symbolic regression discovers a closed-form '
    'equation with only 4 parameters that achieves the best BIC (714, LOO R\u00b2 = 0.665), while '
    'a stacking ensemble achieves R\u00b2 = 0.698 (BIC = 717, LOBO R\u00b2 = 0.707). An expanded '
    'SISSO search with flexible grain-size exponents and EML symbolic regression both fail to '
    'improve upon this equation, validating the constrained search. '
    'However, external validation on 82 independent data points reveals a singularity in the '
    'EN_var/\u03b4_\u03bc term; a robust variant (\u03c3\u00b2_\u03c7 \u2212 \u03a6_VLC replaces '
    '\u03c3\u00b2_\u03c7/\u03b4_\u03bc, LOO R\u00b2 = 0.609) achieves dramatically better out-of-sample '
    'performance (RMSE = 163 vs 421 MPa). '
    '(5) SHAP composition\u2013grain-size interactions (Mn\u00d7d\u207b\xb9, V\u00d7d\u207b\xb9) are '
    'best interpreted as proxies for composition-dependent \u03c3\u2080 rather than composition-dependent '
    'k_HP (see Section 7.3). '
    '(6) The effective Tabor factor C_eff = HV(MPa)/YS = 5.13 \u00b1 1.36 significantly exceeds '
    'the classical value of 3; within Tabor\u2019s H_V \u2248 3\u00b7\u03c3_f(\u03b5_r) generalization (\u03b5_r \u2248 0.08), '
    'this corresponds to an effective early-strain hardening exponent n_eff \u2248 0.15. HV '
    'follows Hall\u2013Petch weakly (R\u00b2 = 0.14) and composition-dependent H\u2080 models do not improve '
    'predictions, unlike the YS case. '
    '(7) HV is a reliable ranking proxy for YS only when composition is held within a narrow '
    'window: global Spearman \u03c1 = 0.46 vs within-batch 0.70\u20130.95 (Simpson\u2019s paradox), and the '
    'partial correlation \u03c1(HV, YS | d) = 0.24 < \u03c1_global shows that grain size is a confounder '
    'while composition (V especially) is the primary scrambler.'
)
doc.add_page_break()

# ============================================================
# 2. DATASET OVERVIEW
# ============================================================
add_heading('2. Dataset Overview', level=1)

add_heading('2.1 Source and Composition', level=2)
add_para(
    'The dataset comprises 94 FCC HEAs fabricated and characterized across six experimental batches '
    '(BBA, BBB, BBC, CBA, CBB, CBC). Each alloy was cold-worked (50\u201360%), recrystallized at '
    '600\u20131000\u00b0C with hold times of 0.5\u20133 h, and characterized for Vickers hardness (HV) '
    'and tensile yield strength (YS). After removing one alloy with missing YS, 93 alloys constitute '
    'the modeling dataset. Grain sizes, measured by EBSD, range from 15 to 212 \u03bcm.'
)

add_table(
    ['Property', 'Min', 'Max', 'Mean', 'Std Dev'],
    [
        ['YS (MPa)', '152', '544', '269', '82'],
        ['HV', '58', '228', '133', '33'],
        ['Grain Size (\u03bcm)', '15', '212', '58', '43'],
        ['Cold Work (%)', '50', '60', '57', '5'],
        ['Recryst. T (\u00b0C)', '600', '1000', '878', '116'],
        ['Hold Time (h)', '0.5', '3.0', '1.0', '0.7'],
    ]
)

add_table(
    ['Element', 'Min (at%)', 'Max (at%)', 'Role in strengthening'],
    [
        ['Al', '0', '4', 'Large misfit (143 pm), FCC destabilizer'],
        ['Co', '0', '52', 'FCC stabilizer, moderate misfit'],
        ['Cr', '0', '20', 'SRO promoter, moderate misfit (128 pm)'],
        ['Cu', '0', '24', 'GB segregant, immiscible with Fe/Cr'],
        ['Fe', '0', '32', 'Base element, magnetic effects'],
        ['Mn', '0', '24', 'SFE modifier, GB segregant'],
        ['Ni', '16', '72', 'Dominant FCC stabilizer (always \u226516%)'],
        ['V', '0', '24', 'Large misfit (134 pm), strongest empirical strengthening effect'],
    ]
)

add_heading('2.2 Feature Engineering', level=2)
add_para(
    'Starting from 8 element fractions, 3 processing parameters, and grain size, we computed 15 '
    'physics-based descriptors and organized features into four hierarchical tiers:'
)
add_table(
    ['Tier', 'Features', 'Count', 'Description'],
    [
        ['CORE', 'Fractions + d\u207b\xb9\u2082 + processing', '12', 'Minimal feature set'],
        ['PHYSICS', 'CORE + descriptors', '27', 'Adds \u03b4, VEC, \u0394H, \u0394S, \u03a9, etc.'],
        ['INTERACTIONS', 'PHYSICS + cross-terms', '46', 'Element\u00d7d\u207b\xb9\u2082 interactions'],
        ['INTERACTIONS_ALT', 'INTERACTIONS + alt-GS', '64', 'Adds d\u207b\xb9, d\u207b\xb9\u2033, etc.'],
        ['COMPACT', 'MI-selected subset', '20', 'Mutual information top-20'],
    ]
)

add_heading('2.3 Notation', level=2)
add_table(
    ['Symbol', 'Definition', 'Units'],
    [
        ['\u03c3_y (YS)', 'Yield strength', 'MPa'],
        ['\u03c3\u2080', 'Friction stress (Hall\u2013Petch intercept)', 'MPa'],
        ['k_HP', 'Hall\u2013Petch coefficient', 'MPa\u00b7\u03bcm\xb9\u2082'],
        ['d (GrainSize)', 'Mean grain size (EBSD)', '\u03bcm'],
        ['x\u1d62', 'Atomic fraction of element i', '\u2014'],
        ['HV', 'Vickers hardness number', 'kgf/mm\u00b2'],
        ['\u03b4', 'Atomic size mismatch', '\u2014'],
        ['\u03a6_VLC', 'VLC misfit variance parameter', '\u2014'],
    ]
)

add_heading('2.4 HEA Descriptors', level=2)
add_table(
    ['Descriptor', 'Definition'],
    [
        ['\u03b4', 'Atomic size mismatch: \u221a(\u2211c_i(1\u2212r_i/r\u0304)\u00b2)'],
        ['VEC', 'Valence electron concentration: \u2211c_i\u00b7VEC_i'],
        ['\u0394H_mix', 'Mixing enthalpy: \u2211 4\u0394H_ij\u00b7c_i\u00b7c_j'],
        ['\u0394S_mix', 'Mixing entropy: \u2212R\u2211c_i\u00b7ln(c_i)'],
        ['\u03a9', 'Stability: T_m\u00b7\u0394S_mix/|\u0394H_mix|'],
        ['\u03bc\u0304', 'Rule-of-mixtures shear modulus'],
        ['\u0394\u03c7', 'Electronegativity mismatch'],
        ['\u03c3_VLC', 'VLC SSS prediction (T=300 K)'],
        ['\u03c3_Labusch', 'Labusch SSS estimate'],
        ['\u03c3_TC', 'Toda-Caraballo SSS estimate'],
        ['\u03b4_Yang', 'Yang lattice distortion parameter'],
        ['\u03a6_VLC', 'VLC misfit variance \u2211c_n\u03b4V_n\u00b2'],
        ['\u03b5_L', 'Combined Labusch strain parameter'],
    ]
)
doc.add_page_break()

# ============================================================
# 3. EXPLORATORY DATA ANALYSIS
# ============================================================
add_heading('3. Exploratory Data Analysis', level=1)

add_heading('3.1 Correlation Structure', level=2)
add_para(
    'Pearson correlation analysis identifies the top correlates with yield strength: d\u207b\xb9\u2082 '
    '(r = 0.66), V content (r = 0.51), and the VLC misfit parameter \u03a6_VLC (r = 0.45). Several '
    'descriptors exhibit high mutual correlation (e.g., \u03b4 and \u03a6_VLC, r > 0.9), motivating '
    'regularized regression and tree-based methods.'
)
add_figure('01_correlation_matrix.png',
           'Figure 1. Pearson correlation matrix for all features and targets. Strong correlations among '
           'SSS descriptors motivate regularization.')

add_heading('3.2 Baseline Hall\u2013Petch', level=2)
add_para(
    'A simple Hall\u2013Petch regression of YS against d\u207b\xb9\u2082 across all 93 alloys yields '
    'Train R\u00b2 = 0.43 (LOO R\u00b2 = 0.41), confirming that grain-size dependence alone accounts for less than half the '
    'observed variance. The residual scatter is strongly correlated with composition, particularly V '
    'content: alloys containing V exhibit systematically higher yield strengths at comparable grain sizes.'
)
add_figure('02_hall_petch.png',
           'Figure 2. Yield strength vs d\u207b\xb9\u2082 for all 93 FCC HEAs. The global regression '
           '(Train R\u00b2 = 0.43, LOO R\u00b2 = 0.41) shows substantial residual variance attributable to composition.')

add_heading('3.3 Composition Effects', level=2)
add_para(
    'Yield strength varies strongly with composition. V exhibits the strongest positive effect on YS, '
    'consistent with its large atomic-size mismatch (134 pm vs \u0101\u0304 \u2248 127 pm). Al and Cu also '
    'show positive associations, while Ni content (which ranges from 16\u201372 at%) shows a weak '
    'negative effect. These composition effects are batch-dependent, reflecting the confounding of '
    'composition with processing in the experimental design.'
)
add_figure('03_YS_vs_composition.png',
           'Figure 3. Yield strength dependencies on individual element concentrations.')
add_figure('06_YS_vs_descriptors.png',
           'Figure 4. Yield strength vs heuristic descriptors (\u03b4, VEC, \u0394S_mix, \u03a9, etc.).')

add_heading('3.4 Processing Effects', level=2)
add_para(
    'Recrystallization temperature is the dominant processing parameter, with higher temperatures '
    'producing larger grains and lower yield strengths. Cold work percentage varies narrowly (50\u201360%) '
    'and contributes minimal variance after recrystallization. Hold time effects are secondary.'
)
add_figure('05_processing_effects.png',
           'Figure 5. Effects of cold work, recrystallization temperature, and hold time on YS and HV.')

add_heading('3.5 Batch-Specific Hall\u2013Petch', level=2)
add_para(
    'Batch-specific Hall\u2013Petch regressions reveal substantial variation in both \u03c3\u2080 and '
    'k_HP across the six experimental campaigns. This variation motivates the LOBO evaluation protocol '
    'and confirms that composition modulates both the friction stress and the grain-boundary '
    'strengthening coefficient.'
)
add_figure('09_batch_hall_petch.png',
           'Figure 6. Batch-specific Hall\u2013Petch regressions showing variation in \u03c3\u2080 and k_HP.')

add_heading('3.6 Random Forest Feature Importance (Preliminary)', level=2)
add_para(
    'A preliminary random forest analysis identifies the most predictive features for YS: d\u207b\xb9\u2082 '
    '(dominant), V fraction, \u03b4 (lattice distortion), \u03a9 (stability parameter), and Mn fraction.'
)
add_figure('07_feature_importance_YS.png',
           'Figure 7. Random forest feature importance for yield strength prediction.')

add_heading('3.7 Pre-modeling Diagnostics', level=2)
add_para(
    'Three diagnostics motivated by data-design considerations were performed before any '
    'targeted regression: (i) composition × grain-size confounding, (ii) batch × '
    'composition coverage in PCA space, and (iii) the variance floor implied by '
    'within-composition replicates and per-alloy measurement uncertainties.'
)

add_bold_then_text(
    'Composition × grain-size confounding. ',
    'Univariate Pearson correlations between each elemental fraction and grain size identify '
    'Cr (r = +0.38, p = 2×10⁻⁴) and Al (r = −0.32, p = 2×10⁻³) '
    'as the strongest correlates; V (r = −0.24) and Cu (r = −0.21) are weaker but '
    'significant. The joint OLS regression R²(d⁻\xb9₂ ~ 8 element fractions) = 0.40, '
    'rising to 0.54 when the three processing variables are added; ~46% of d⁻\xb9₂ '
    'variance therefore remains independent of composition and processing. Restricting attention '
    'to the largest grain sizes (top decile, d ≥ 110 μm), every alloy has Al = 0 and '
    'Cu = 0, all sourced from batches CBA and CBC, so the apparent suppression of coarsening by '
    'Al and Cu is partly batch-design confounded.'
)
add_figure('71_comp_gs_confounding.png',
           'Figure 8. Composition × grain-size confounding: grain size vs. each element '
           '(at.%), batch-coded. Cr and Al are the strongest correlates; the largest grain sizes '
           'occur exclusively in Al-free, Cu-free alloys from CBA/CBC.')

add_bold_then_text(
    'Batch × composition coverage. ',
    'Principal-component analysis of the 8 elemental fractions across the 6 batches gives '
    'effective dimensionality 6 (cumulative variance ≥ 0.95). The mean fraction of one '
    'batch\'s points falling inside the convex hull of another (PC1–PC2 plane) is 0.21; '
    'per-batch held-out coverage when treated as the LOBO target ranges from 0.29 (CBC, most '
    'isolated) to 0.87 (CBA, most reachable). The aggregate LOBO R² therefore averages '
    'over qualitatively different tests: holding out CBC or BBB tests genuine compositional '
    'extrapolation, whereas holding out CBA tests resampling within an already-covered region. '
    'Per-batch LOBO breakdowns would be more interpretable than the aggregate value.'
)
add_figure('72_batch_pca_coverage.png',
           'Figure 9. Batch × composition coverage: (a) PCA scatter colored by batch, '
           '(b) convex hulls per batch, (c) heatmap of pairwise hull-containment fractions.')

add_bold_then_text(
    'Variance floor and R² ceilings. ',
    'Of the 81 unique compositions, 72 are singletons and 9 have ≥ 2 alloys (21 alloys '
    'total, 22.6% pseudo-replication). A perfect composition-mean predictor would leave only '
    'the within-group YS variance as residual, implying a composition-only R² ceiling of '
    'approximately 0.86. The reported per-alloy SD_YS values (rms = 15.3 MPa over 80 alloys) '
    'imply ⟨SD_YS²⟩ = 234 MPa² and a noise-limited R² ceiling of '
    'approximately 0.97 for any model that uses composition and grain size. XGBoost (LOO '
    'R² = 0.729) operates at ~76% of this noise-limited ceiling, M3 at ~68%, and SISSO '
    'Robust at ~65%; roughly one-quarter of the explainable variance remains unexplained.'
)
add_figure('73_pseudoreplicate_variance.png',
           'Figure 10. Variance decomposition: (a) histogram of replicate-group sizes, '
           '(b) within-group YS SD vs. ΔGrainSize within group, '
           '(c) total / within-comp / measurement variance bars with implied R² ceilings.')

doc.add_page_break()

# ============================================================
# 4. SOLID-SOLUTION STRENGTHENING MODELS
# ============================================================
add_heading('4. Solid-Solution Strengthening Models', level=1)

add_heading('4.1 Model Descriptions', level=2)

add_bold_then_text('Varvenne\u2013Leyson\u2013Curtin (VLC): ',
    'The VLC model (Varvenne et al., Acta Mater. 2016) computes SSS from atomic volume misfit '
    'variance: \u03c4_y0 = 0.051\u00b7\u03b1\u207b\xb9\u2033\u00b7f\u2081\u00b2\u2033\u00b7\u03bc\u0304'
    '\u00b7Q\u00b2\u2033, where Q = \u2211c_n(\u0394V_n/V\u0304)\u00b2. The energy barrier for '
    'dislocation glide is \u0394E_b = 0.274\u00b7\u03b1\xb9\u2033\u00b7f\u2081\xb9\u2033\u00b7\u03bc\u0304'
    '\u00b7b\u00b3\u00b7Q\xb9\u2033. Temperature correction: \u03c3(T) = M\u00b7\u03c4_y0\u00b7'
    '[1 \u2212 (kT/\u0394E_b)\u00b2\u2033], M = 3.06.')

add_bold_then_text('Labusch model: ',
    'Combined size + modulus misfits: \u03c3_Labusch = \u03bc\u0304\u00b7\u03b5_L\u2074\u2033'
    '\u00b7c_eff\xb9\u2033, where \u03b5_L includes both radius and shear modulus mismatch with '
    '\u03b1_L = 16 for edge dislocations.')

add_bold_then_text('Toda-Caraballo (TC) model: ',
    '\u03c3_TC = M\u00b7\u03bc\u0304\u00b7\u03b4_Yang\u2074\u2033\u00b7c_eff\xb9\u2033, '
    'where \u03b4_Yang = \u221a(\u2211c_i(\u0394r_i/r\u0304)\u00b2) is the Yang lattice distortion.')

add_heading('4.2 Bug Fix in VLC Implementation', level=2)
add_para(
    'During analysis, a critical bug was identified in the original VLC energy barrier formula. The '
    'implementation used \u0394E_b = 0.274\u00b7\u0393\xb9\u2033\u00b7(\u03bc\u00b7b\u2075\u00b7f\u2081'
    '\u00b7Q)\u00b2\u2033, which produces ~10\u207b\u00b3\xb9 J (essentially zero), making all 300 K '
    'predictions = 0. The corrected formula \u0394E_b = 0.274\u00b7\u03b1\xb9\u2033\u00b7f\u2081\xb9\u2033'
    '\u00b7\u03bc\u00b7b\u00b3\u00b7Q\xb9\u2033 gives physically meaningful values (0.058 eV for '
    'CoCrFeMnNi, kT/\u0394E_b = 0.44 at 300 K). All results below use the corrected implementation.'
)

add_heading('4.3 Quantitative Results', level=2)
add_para(
    'None of the three models produces quantitatively accurate absolute predictions of experimental '
    'yield strength. The Labusch model\'s large Pred/Exp ratio (28.4\u00d7, after correcting the '
    'concentration exponent from c\u02b8\u00b9\u0301\u00b3 to c\u02b8\u00b2\u0301\u00b3 per Labusch 1970) reflects the '
    'absence of the 1/Z numerical prefactor in our empirical HEA extension of the dilute formula. '
    'Toda-Caraballo, now implemented faithfully with the Gypen-Deruyttere superposition rule '
    '(Eq. 14 of their 2015 paper) and Z = 1/180 from Labusch\'s original derivation (no fitting), '
    'overpredicts CoCrFeMnNi by a factor of ~3.6 (446 vs. 125 MPa experimental). The ranking '
    'correlations cluster in a narrow band (\u03c1 = 0.54-0.57 for the three models) and remain '
    'the more useful metric.'
)

add_table(
    ['SSS Model', 'Mean (MPa)', 'Pred/Exp Ratio', 'Pearson r', 'Raw R\u00b2', '+HP LOO R\u00b2'],
    [
        ['VLC (300 K)', '342', '1.27', '0.261', '\u22123.4', '0.396'],
        ['Labusch', '7,634', '28.4', '0.553', '\u22128,759', '0.469'],
        ['Toda-Caraballo', '978', '3.64', '0.355', '\u221296.9', '0.399'],
        ['HP only', '\u2014', '\u2014', '\u2014', '\u2014', '0.406'],
        ['Expt. YS', '269', '1.0', '\u2014', '\u2014', '\u2014'],
    ]
)

add_heading('4.4 Standalone vs Combined Performance (LOO Ridge)', level=2)
add_table(
    ['Model', 'LOO R\u00b2', 'RMSE (MPa)', 'MAE (MPa)'],
    [
        ['Hall\u2013Petch only', '0.406', '62.5', '47.1'],
        ['VLC(300K) only', '0.030', '79.9', '63.4'],
        ['Labusch only', '0.274', '69.1', '53.0'],
        ['TC only', '0.091', '77.3', '61.2'],
        ['HP + VLC', '0.397', '63.0', '47.2'],
        ['HP + Labusch', '0.469', '59.1', '42.0'],
        ['HP + TC', '0.399', '62.9', '46.7'],
        ['Comp + HP', '0.654', '47.7', '30.5'],
        ['Comp + HP + VLC', '0.654', '47.7', '30.5'],
        ['Comp + HP + TC', '0.653', '47.8', '30.5'],
        ['Comp + HP + Lab', '0.655', '47.7', '30.2'],
        ['Comp + HP + All SSS + Proc', '0.589', '52.0', '32.4'],
    ]
)

add_para(
    'Notably, including all SSS descriptors and processing parameters (Comp + HP + All SSS + Proc) '
    'degrades LOO R\u00b2 from 0.652 to 0.588 relative to the composition-only model '
    '(Comp + HP). This degradation likely reflects multicollinearity between SSS descriptors and '
    'their constituent composition fractions, inflating LOO prediction variance, rather than '
    'indicating that SSS physics is incorrect. The Vegard\u2019s-law SSS proxies carry negligible '
    'nonlinear information beyond what composition fractions already provide.'
)

add_heading('4.5 Partial Correlation Analysis', level=2)
add_para(
    'After regressing out composition fractions, the residual correlations between SSS predictions and '
    'YS are negligible: VLC r_partial = \u22120.098, Labusch r_partial = \u22120.069, TC r_partial = '
    '\u22120.071. This indicates that these SSS models, as implemented with Vegard\'s-law inputs, carry '
    'negligible unique information beyond what elemental fractions already provide. The SSS '
    'transformations are deterministic functions of composition and offer minimal additional predictive '
    'power when composition is directly available. This does not necessarily invalidate the underlying '
    'physics; rather, the Vegard\'s-law approximation for misfit volumes may obscure the actual '
    'strengthening contributions.'
)

add_heading('4.6 Vegard\'s Law: The Root Cause', level=2)
add_para(
    'For equimolar CoCrFeMnNi, our Vegard\'s-law VLC predicts \u03c3_VLC(300K) \u2248 30 MPa, '
    'substantially below the ~125\u2013200 MPa obtained by Varvenne et al. and Moitzi et al. using '
    'DFT-derived misfit volumes. However, for V-rich and Al-rich compositions, the same model '
    'overpredicts by 4\u20135\u00d7 because Vegard\'s law linearly extrapolates lattice parameters '
    'from pure-element FCC structures that V and Al do not naturally adopt. This sign-dependent error '
    'cannot be corrected by simple rescaling. The VLC theory itself is sound; the input approximation '
    '(Vegard\'s law) is the bottleneck.'
)

add_figure('13_vlc_sss_analysis.png',
           'Figure 11. VLC SSS analysis: (a) VLC predictions vs experimental YS by batch, (b) residual '
           'YS after HP removal vs VLC, (c) best linear superposition parity, (d) VLC vs YS colored by V content.')
add_figure('14_strengthening_decomposition.png',
           'Figure 12. Decomposed strengthening contributions: \u03c3\u2080 + \u0394\u03c3_SSS(VLC) + '
           '\u0394\u03c3_HP vs experimental YS, sorted by increasing experimental YS.')
doc.add_page_break()

# ============================================================
# 5. GRAIN-SIZE SCALING LAWS
# ============================================================
add_heading('5. Grain-Size Scaling Laws', level=1)

add_heading('5.1 Models Tested', level=2)
add_para(
    'We compare nine grain-size scaling laws of the general form \u03c3_y = \u03c3\u2080 + k\u00b7f(d):')
add_table(
    ['#', 'Scaling f(d)', 'Parameters', 'Origin'],
    [
        ['1', 'd\u207b\xb9\u2082', '2', 'Hall\u2013Petch (1951/1953)'],
        ['2', 'd\u207b\xb9', '2', 'Dunstan\u2013Bushby (2014)'],
        ['3', 'd\u207b\xb9\u2033', '2', 'Baldwin (1958)'],
        ['4', 'd\u207b\u00b2\u2033', '2', 'Intermediate exponent'],
        ['5', 'ln(d)/d', '2', 'Critical thickness model'],
        ['6', 'ln(d)', '2', 'Logarithmic'],
        ['7', '1/\u221ad + 1/d', '3', 'Composite (two-term)'],
        ['8', 'd\u207b\xb9 + d\u207b\u00b2', '3', 'Taylor expansion'],
        ['9', 'd\u207b\u207f_opt', '3', 'Optimized exponent'],
    ]
)

add_heading('5.2 Results', level=2)
add_table(
    ['Scaling f(d)', 'k', 'Train R\u00b2', 'LOO R\u00b2', 'LOO RMSE', 'BIC', '\u0394BIC'],
    [
        ['d\u207b\xb9\u2082 [HP]', '2', '0.429', '0.406', '62.5', '774.5', '0.0'],
        ['ln(d)/d', '2', '0.429', '0.406', '62.5', '774.5', '0.0'],
        ['d\u207b\u00b2\u2033', '2', '0.429', '0.405', '62.6', '774.6', '0.1'],
        ['d\u207b\xb9\u2033 [Baldwin]', '2', '0.427', '0.403', '62.6', '774.9', '0.4'],
        ['d\u207b\xb9 [D\u2013B]', '2', '0.419', '0.395', '63.1', '776.1', '1.6'],
        ['ln(d)', '2', '0.413', '0.388', '63.4', '777.2', '2.7'],
        ['d\u207b\u2070\u00b7\u2075\u2074\u2078 [Opt]', '3', '0.430', '0.406', '62.5', '779.0', '4.5'],
        ['1/\u221ad + 1/d', '3', '0.430', '0.395', '63.1', '779.0', '4.5'],
        ['d\u207b\xb9 + d\u207b\u00b2', '3', '0.429', '0.387', '63.5', '779.1', '4.6'],
    ]
)

add_para(
    'Four scaling laws (d\u207b\xb9\u2082, ln(d)/d, d\u207b\u00b2\u2033, d\u207b\xb9\u2033) are '
    'statistically indistinguishable at the \u0394BIC < 2 threshold. This indistinguishability reflects '
    'the limited statistical power of the grain-size-only model (R\u00b2 \u2248 0.43), where composition-driven '
    'scatter dominates grain-size effects, rather than physical equivalence of the scaling laws. '
    'The optimized exponent is n_opt = 0.548, close to the classical value of 0.5. The d\u207b\xb9 '
    '(Dunstan\u2013Bushby) and ln(d) scalings perform distinctly worse (\u0394BIC > 1.6).'
)

add_figure('25_scaling_comparison.png',
           'Figure 13. Grain-size scaling law comparison: (a) fits of top scaling laws, (b) \u0394BIC '
           'comparison (green = strong support, orange = moderate, red = weak), (c) LOO R\u00b2 vs BIC.')

# ============================================================
# 6. BAYESIAN GRAIN-SIZE SCALING ANALYSIS
# ============================================================
doc.add_page_break()
add_heading('6. Bayesian Grain-Size Scaling Analysis', level=1)

add_heading('6.1 Bayesian Model Comparison (PSIS-LOO)', level=2)
add_para(
    'To complement the frequentist \u0394BIC analysis, we conduct a full Bayesian analysis using '
    'PyMC (v5.28) with Markov Chain Monte Carlo (MCMC) sampling. Each scaling law is fitted as a '
    'Bayesian linear regression: \u03c3_y = \u03c3\u2080 + k\u00b7f(d) + \u03b5, where '
    '\u03b5 ~ N(0, \u03c3\u00b2). Weakly informative priors ensure a fair comparison across models: '
    '\u03c3\u2080 ~ N(200, 200), k ~ N(0, 1000), \u03c3 ~ HalfCauchy(50). For the optimized-exponent '
    'model, n ~ Uniform(0.1, 2.0). All models are sampled with 4,000 draws, 2,000 tuning steps, '
    'and 2 chains (below the recommended minimum of 4; additional chains would strengthen convergence '
    'diagnostics, particularly for the nonlinear exponent model). Convergence diagnostics confirm '
    'R\u0302 < 1.002 and ESS > 1,400 for all parameters.'
)

add_para(
    'Model comparison uses Pareto-smoothed importance sampling leave-one-out cross-validation '
    '(PSIS-LOO; Vehtari et al., 2017), the gold-standard Bayesian model comparison method. '
    'PSIS-LOO estimates the expected log pointwise predictive density (elpd_loo), which measures '
    'predictive accuracy on held-out data. Stacking weights (Yao et al., 2018) provide '
    'Bayesian model probabilities optimized for out-of-sample prediction.'
)

add_table(
    ['Model', 'Rank', 'elpd_loo', '\u0394LOO', 'p_loo', 'Weight'],
    [
        ['Hall\u2013Petch (d\u207b\xb9\u2082)', '1', '\u2212518.5', '0.0', '4.2', '1.000'],
        ['Composite (d\u207b\xb9\u2082 + d\u207b\xb9)', '2', '\u2212518.5', '0.1', '4.2', '0.000'],
        ['Crit. thickness (ln d/d)', '3', '\u2212518.5', '0.1', '4.3', '0.000'],
        ['Intermediate (d\u207b\u00b2\u2033)', '4', '\u2212518.7', '0.2', '4.4', '0.000'],
        ['Baldwin (d\u207b\xb9\u2033)', '5', '\u2212518.7', '0.2', '4.2', '0.000'],
        ['Optimized exponent (d\u207b\u207f)', '6', '\u2212518.8', '0.3', '4.5', '0.000'],
        ['Logarithmic (ln d)', '7', '\u2212519.7', '1.2', '3.9', '0.000'],
        ['Dunstan\u2013Bushby (d\u207b\xb9)', '8', '\u2212520.3', '1.8', '4.3', '0.000'],
    ]
)

add_para(
    'The PSIS-LOO results are broadly consistent with the \u0394BIC analysis. Hall\u2013Petch ranks '
    'first, but the top six models are effectively indistinguishable (\u0394LOO < 0.3). Stacking '
    'weights concentrate entirely on Hall\u2013Petch (weight = 1.0), indicating that when forced to '
    'choose a single model for point prediction, the classical d\u207b\xb9\u2082 law is preferred. '
    'However, the negligible \u0394LOO values mean that the data cannot discriminate between exponents '
    'in the range ~0.3\u20130.7. The logarithmic (\u0394LOO = 1.2) and Dunstan\u2013Bushby d\u207b\xb9 '
    '(\u0394LOO = 1.8) models show mild disfavor but remain within the range of plausibility.'
)

add_figure('30_bayesian_model_comparison.png',
           'Figure 14. Bayesian model comparison: PSIS-LOO expected log predictive density '
           'differences (\u0394elpd) with standard errors for all 8 grain-size scaling laws.')

add_heading('6.2 Parameter Posteriors', level=2)
add_para(
    'Posterior distributions for the key parameters (\u03c3\u2080 and k) reveal consistent '
    'interpretations across models. The friction stress \u03c3\u2080 ranges from 11 MPa (Baldwin) to '
    '570 MPa (Logarithmic), reflecting the different functional forms rather than physical differences. '
    'The Hall\u2013Petch posterior yields \u03c3\u2080 = 92 \u00b1 22 MPa and k_HP = 1149 \u00b1 140 '
    'MPa\u00b7\u03bcm\xb9\u2082, with 94% highest density intervals (HDI) of [51, 135] and '
    '[887, 1419] respectively. The noise parameter \u03c3 \u2248 62.5 MPa is consistent across all models.'
)

add_table(
    ['Model', '\u03c3\u2080 (MPa)', '94% HDI', 'k', '94% HDI', '\u03c3 (MPa)'],
    [
        ['Hall\u2013Petch (d\u207b\xb9\u2082)', '91.9', '[51, 135]', '1149', '[887, 1419]', '62.5'],
        ['Dunstan\u2013Bushby (d\u207b\xb9)', '190.4', '[166, 214]', '3044', '[2206, 3788]', '63.5'],
        ['Baldwin (d\u207b\xb9\u2033)', '10.9', '[\u221246, 72]', '907', '[703, 1111]', '62.5'],
        ['Intermediate (d\u207b\u00b2\u2033)', '135.0', '[102, 169]', '1592', '[1199, 1952]', '62.4'],
        ['Crit. thickness (ln d/d)', '143.9', '[114, 175]', '1385', '[1086, 1716]', '62.4'],
        ['Logarithmic (ln d)', '570.2', '[502, 647]', '\u221278.6', '[\u221297, \u221260]', '63.2'],
    ]
)

add_figure('31_bayesian_posteriors.png',
           'Figure 15. Forest plot of posterior distributions for \u03c3\u2080 and k across all models. '
           'Points show posterior medians; thick bars show 50% HDI; thin bars show 94% HDI.')

add_heading('6.3 Posterior Predictive Checks', level=2)
add_para(
    'Posterior predictive checks (PPCs) assess model adequacy by comparing the observed data '
    'distribution against data simulated from the fitted model. For the top-ranked models, the '
    'posterior predictive distributions capture the overall shape and spread of the observed YS '
    'distribution, though all models show similar limitations in capturing the full heterogeneity '
    'arising from composition effects.'
)

add_figure('32_bayesian_ppc.png',
           'Figure 16. Posterior predictive checks for the top 4 scaling laws. Blue histograms show '
           'simulated data; orange line shows observed YS distribution.')

add_heading('6.4 Optimized Exponent Posterior', level=2)
add_para(
    'The optimized-exponent model samples the scaling exponent n from data. The posterior yields '
    'n = 0.53 \u00b1 0.16 with a 94% HDI of [0.22, 0.82]. The posterior probability that n < 0.5 '
    '(favoring a weaker-than-HP scaling) is 0.43, while P(0.4 < n < 0.6) = 0.43, confirming that '
    'the data is consistent with the classical HP exponent of 0.5 but cannot rule out exponents '
    'between approximately 0.2 and 0.8. This broad posterior reflects the intrinsic difficulty of '
    'distinguishing power-law exponents from data with substantial composition-driven scatter.'
)

add_figure('34_bayesian_exponent.png',
           'Figure 17. Posterior distribution of the grain-size scaling exponent n from the '
           'optimized-exponent model. Dashed lines mark key theoretical values: 0.5 (Hall\u2013Petch), '
           '0.33 (Baldwin), 0.67 (Intermediate), 1.0 (Dunstan\u2013Bushby).')

add_heading('6.5 Bayesian Model Averaging', level=2)
add_para(
    'Bayesian Model Averaging (BMA) combines predictions from all models weighted by their stacking '
    'weights. Because all weight concentrates on Hall\u2013Petch, the BMA prediction is effectively '
    'identical to the HP model. The 94% posterior credible band spans approximately \u00b1125 MPa, '
    'reflecting the large irreducible variance from composition effects not captured by grain-size-only '
    'models.'
)

add_figure('33_bayesian_bma.png',
           'Figure 18. Bayesian Model Averaged prediction (green line) with 94% credible band '
           '(shaded region) vs observed data points.')

add_figure('35_bayesian_weights.png',
           'Figure 19. Stacking weights for Bayesian Model Averaging. Hall\u2013Petch receives '
           'weight = 1.0; all other models receive negligible weight. Note: the degenerate stacking weights '
           '(HP \u2248 1.0) reflect a corner solution of the Bayesian stacking optimization when models have '
           'nearly identical elpd_loo (\u0394LOO < 0.3), and should not be interpreted as strong evidence '
           'favoring HP over the other five top-ranked scaling laws.')

# ============================================================
# 7. COMPOSITION-DEPENDENT HALL-PETCH MODELS
# ============================================================
doc.add_page_break()
add_heading('7. Composition-Dependent Hall\u2013Petch Models', level=1)

add_heading('7.1 Model Hierarchy', level=2)
add_para(
    'The grain-size-only models in Sections 5 and 6 treat \u03c3\u2080 and k as global constants, '
    'yielding LOO R\u00b2 \u2248 0.41 and k_HP \u2248 1149 MPa\u00b7\u03bcm\xb9\u2082. This value is '
    'substantially higher than literature values for FCC HEAs (494\u2013824; see Section 7.4). The '
    'discrepancy arises because the grain-size-only model conflates composition effects with the HP '
    'slope: V-rich alloys are systematically stronger and tend to have finer grains (r(V, d) = \u22120.24), '
    'so the global regression attributes some of V\'s composition-dependent \u03c3\u2080 effect to the '
    'd\u207b\xb9\u2082 slope, inflating k. When composition enters the model (below), k drops to 766, '
    'within the expected range.'
)

add_para(
    'As will be shown in the SHAP analysis (Section 9), there are strong '
    'composition\u2013grain-size interactions (V\u00d7d\u207b\xb9, Mn\u00d7d\u207b\xb9), we next examine '
    'whether allowing \u03c3\u2080 and/or k_HP to depend on composition improves predictions. '
    'We test 13 models of the form \u03c3_y = \u03c3\u2080(comp) + k(comp)\u00b7d\u207b\xb9\u2082, '
    'spanning four groups: (A) composition-dependent \u03c3\u2080 only, (B) composition-dependent k only, '
    '(C) both, and (D) physics descriptors.'
)

add_table(
    ['Model', 'Params', 'LOO R\u00b2', 'RMSE', '\u0394BIC', 'PSIS-LOO rank'],
    [
        ['M0: Baseline HP', '3', '0.406', '62.5', '36.5', '13'],
        ['M1: \u03c3\u2080(V)', '4', '0.605', '51.0', '1.4', '5'],
        ['M3: \u03c3\u2080(all 7 elem)', '10', '0.652', '47.8', '0.0', '2'],
        ['M4: k(V)', '4', '0.598', '51.4', '3.2', '9'],
        ['M6: k(all 7 elem)', '10', '0.622', '49.9', '5.5', '3'],
        ['M7: \u03c3\u2080(V)+k(V)', '5', '0.598', '51.4', '5.9', '4'],
        ['M8: \u03c3\u2080(V,Mn)+k(V,Mn)', '7', '0.597', '51.5', '11.9', '7'],
        ['M10: \u03c3\u2080(all)+k(all)', '17', '0.650', '48.0', '17.3', '1'],
        ['M11: \u03c3\u2080(\u03b4)', '4', '0.456', '59.8', '31.3', '12'],
    ]
)

add_para(
    'Note: M0 (Baseline HP) reports LOO R\u00b2 = 0.406 with 3 parameters (intercept, slope, noise '
    'variance). The Section 5 scaling-law HP analysis uses k = 2 (intercept + slope), so BIC '
    'values differ between the two tables, but the LOO R\u00b2 is identical.'
)

add_para(
    'The single most impactful improvement is making \u03c3\u2080 composition-dependent. Adding just one '
    'parameter\u2014the V fraction\u2014to the friction stress (M1) raises LOO R\u00b2 from 0.406 to '
    '0.605 (+0.199 in R\u00b2, a 33% reduction in unexplained variance). The best BIC model (M3) uses '
    'all 7 non-Ni element fractions in \u03c3\u2080 with a constant k_HP, achieving LOO R\u00b2 = 0.652 '
    '(RMSE = 47.8 MPa). However, M1 (\u03c3\u2080 depends only on V fraction) achieves \u0394BIC = 1.4 '
    'relative to M3, making the two models statistically indistinguishable by the \u0394BIC < 2 '
    'criterion used for scaling laws. The simpler M1 model\u2014in which V content alone captures '
    'two-thirds of the composition effect\u2014represents an equally valid parsimonious alternative, and '
    'the 7-coefficient interpretation of M3 should be understood as contingent on this model choice. '
    'Bayesian stacking assigns weight = 0.90 to M3; note that M10 (\u03c3\u2080(all)+k(all)) achieves the '
    'highest elpd_loo point estimate (\u2212493.5 vs M3\u2019s \u2212494.0), but its PSIS-LOO diagnostics '
    'indicate unreliable individual LOO estimates (Pareto k\u0302 > 0.7), making its elpd potentially '
    'biased. This supports M3 as the preferred specification, though M1 remains a viable alternative.'
)

add_heading('7.2 Fitted Composition-Dependent \u03c3\u2080', level=2)
add_para(
    'The best model (M3) takes the form \u03c3_y = [\u03c3\u2080\u2080 + \u03a3\u03b1\u1d62\u00b7x\u1d62] + '
    'k\u00b7d\u207b\xb9\u2082, where the 7 non-Ni fractions enter linearly into \u03c3\u2080 and the '
    'intercept \u03c3\u2080\u2080 absorbs the Ni (solvent) contribution.'
)

add_table(
    ['Parameter', 'Value', 'Physical meaning'],
    [
        ['\u03c3\u2080\u2080', '230 MPa', 'Intercept (regression constant; not a physical estimate of pure Ni strength)'],
        ['\u03b1_V', '+291 MPa', 'V strongly increases \u03c3\u2080 (large misfit)'],
        ['\u03b1_Mn', '+82 MPa', 'Mn mildly increases \u03c3\u2080'],
        ['\u03b1_Al', '\u2212201 MPa', 'Al decreases \u03c3\u2080 relative to Ni'],
        ['\u03b1_Co', '\u2212187 MPa', 'Co decreases \u03c3\u2080'],
        ['\u03b1_Cr', '\u2212308 MPa', 'Cr decreases \u03c3\u2080'],
        ['\u03b1_Cu', '\u2212334 MPa', 'Cu decreases \u03c3\u2080'],
        ['\u03b1_Fe', '\u2212360 MPa', 'Fe largest negative effect on \u03c3\u2080'],
        ['k_HP', '766 MPa\u00b7\u03bcm\xb9\u2082', 'Global constant'],
    ]
)

add_para(
    'The friction stress \u03c3\u2080(comp) ranges from 84 to 269 MPa across the 93 alloys, with V-rich '
    'compositions exhibiting the highest values. The \u03b1\u1d62 coefficients represent empirical '
    'regression weights capturing the net effect of each element on strength at infinite grain size '
    '(d \u2192 \u221e). This includes but is not limited to solid-solution strengthening; the '
    'coefficients may also absorb short-range ordering, stacking-fault energy modifications, and '
    'other composition-dependent effects. The individual \u03b1\u1d62 values should not be '
    'interpreted as physics-based SSS contributions. Composition modulates the baseline lattice '
    'resistance, while the grain-boundary strengthening efficiency (k_HP) shows no statistically '
    'significant composition dependence in this dataset (Section 7.3).'
)
add_para(
    'The dominant role of V (\u03b1_V = +291 MPa, the largest positive coefficient) is consistent '
    'with recent theoretical predictions by Yin, Maresca, and Curtin (Acta Materialia 188, 2020), '
    'who demonstrated that V is an optimal element for solid-solution strengthening in both FCC and '
    'BCC HEAs. The origin is V\u2019s anomalously large atomic volume in the FCC matrix '
    '(134 pm vs. \u0101 \u2248 127 pm for Co-Cr-Fe-Mn-Ni), which produces the largest misfit volume '
    'among the constituent elements. Their parameter-free Varvenne\u2013Curtin theory predicts that '
    'V additions at \u223c25 at.% maximize SSS in FCC Co-Cr-Fe-Mn-Ni-V alloys. Our empirical M1 model '
    'confirms this picture: adding V fraction alone as a single composition descriptor to the '
    'Hall\u2013Petch equation raises LOO R\u00b2 from 0.406 to 0.605, capturing approximately '
    'two-thirds of the total composition effect on \u03c3\u2080. This convergence between '
    'first-principles theory and data-driven regression provides suggestive, though not conclusive, '
    'agreement\u2014the '
    'Varvenne\u2013Curtin mechanism rationalizes why V dominates our empirical coefficients, while '
    'our experimental data confirms the predicted magnitude of V\u2019s strengthening effect across '
    'a broad, non-equimolar composition space.'
)
add_figure('49_misfit_vs_coefficients.png',
           'Figure 20. Varvenne\u2013Curtin misfit theory vs empirical M3 coefficients: '
           '(a) Goldschmidt radii for all 8 elements with dataset-mean \u0101 = 126.1 pm. '
           'V (+6.2%) and Al (+13.4%) exhibit the largest positive misfits. '
           '(b) Empirical \u03b1\u1d62 coefficients from M3 (relative to Ni reference). '
           'V dominates with +291 MPa. '
           '(c) Absolute volume misfit |\u0394V/V\u0305| vs \u03b1\u1d62, showing that V\u2019s '
           'anomalously large misfit volume rationalizes its dominant strengthening effect, '
           'consistent with the parameter-free predictions of Yin, Maresca & Curtin (2020).')

add_figure('36_comp_hp_model_comparison.png',
           'Figure 21. Composition-dependent HP model comparison: (a) LOO R\u00b2, (b) \u0394BIC, '
           '(c) PSIS-LOO. Colors indicate model group: green = \u03c3\u2080 only, blue = k only, '
           'orange = both, purple = physics descriptors.')

add_figure('37_comp_hp_parity.png',
           'Figure 22. LOO parity plots for the top 4 composition-dependent HP models, '
           'colored by V content.')

add_figure('48_M1_model_detail.png',
           'Figure 23. The minimal M1 model (\u03c3_y = \u03c3\u2080\u2080 + \u03b1_V\u00b7x_V + '
           'k\u00b7d\u207b\xb9\u2044\u00b2): (a) LOO parity plot distinguishing V-containing (red '
           'diamonds, n = 84) from V-free alloys (blue circles, n = 9); (b) model surface showing '
           'parallel Hall\u2013Petch lines at V = 0\u201333%, with data colored by V fraction; '
           '(c) violin plots of |LOO residuals| for M0 (baseline HP, 2 params), M1 (+V, 3 params), '
           'and M3 (all elements, 9 params). Median absolute error drops from 38 \u2192 28 \u2192 '
           '19 MPa, indicating that V alone captures approximately two-thirds of the total '
           'composition effect.')

add_heading('7.3 Is k_HP Composition-Dependent?', level=2)
add_para(
    'A refined two-stage analysis tests whether k_HP itself varies with composition. First, \u03c3\u2080(comp) '
    'is estimated from M3 for each alloy. Then, the effective Hall\u2013Petch coefficient is computed as '
    'k_eff = (YS \u2212 \u03c3\u2080(comp)) / d\u207b\xb9\u2082 for each alloy. If k_HP is truly '
    'composition-dependent, k_eff should correlate with elemental fractions. Note that this procedure is '
    'structurally biased toward finding no composition dependence in k_eff: because M3 assumes constant k, '
    'any true composition dependence of k_HP may be partially absorbed into the \u03c3\u2080(comp) estimates, '
    'attenuating the residual signal in the second stage.'
)

add_para(
    'No element shows a statistically significant correlation with k_eff in this dataset '
    '(all |r| < 0.06, all p > 0.6). A multivariate regression of k_eff on all 7 elements yields '
    'R\u00b2 = 0.006 (F-test p = 0.999), indicating that composition explains less than 1% of k_eff '
    'variance. Bayesian analysis corroborates this: the posterior coefficients \u03b2_V = +230 \u00b1 870 '
    'and \u03b2_Mn = \u2212306 \u00b1 913 have uncertainties 3\u20134\u00d7 their means, with posterior probabilities '
    'barely above chance (P(\u03b2_V > 0) = 0.60). However, the limited sample size (N = 93) and '
    'substantial grain-size measurement uncertainty (within-alloy CV \u2248 0.74) may mask weak '
    'composition effects on k_HP.'
)

add_table(
    ['Statistic', 'Value', 'Interpretation'],
    [
        ['k_eff mean', '764 \u00b1 274 MPa\u00b7\u03bcm\xb9\u2082', 'Consistent with global k = 766'],
        ['k_eff CV', '36%', 'Large scatter from measurement uncertainty, grain-size distribution effects, and texture; not explained by composition'],
        ['R\u00b2 (k_eff ~ comp)', '0.006', 'Composition explains <1% of k_eff variance'],
        ['F-test p-value', '0.999', 'No significance whatsoever'],
        ['\u0394LOO: k(V,Mn) vs k=const', '0.5', 'Negligible improvement'],
    ]
)

add_para(
    'This resolves an apparent contradiction with the SHAP analysis. SHAP identified V\u00d7d\u207b\xb9 '
    'and Mn\u00d7d\u207b\xb9 interactions as important features, which could be interpreted as evidence '
    'for composition-dependent k_HP. However, the refined analysis is consistent with the interpretation '
    'that these interaction terms are proxies for composition-dependent \u03c3\u2080 rather than '
    'composition-dependent k. Once \u03c3\u2080(comp) properly absorbs elemental effects, k_HP \u2248 '
    '766 MPa\u00b7\u03bcm\xb9\u2082 is consistent with composition-independent behavior across the 93 '
    'alloys in this dataset, though the two-stage procedure\u2019s structural bias and the 36% '
    'unexplained CV in k_eff leave open the possibility of weak composition dependence below the '
    'detection threshold.'
)

add_para(
    'A within-composition replicate test was also attempted: the 9 composition-replicate groups '
    '(21 alloys at shared elemental compositions but varied grain sizes) should, in principle, '
    'support direct local k_HP fits. Inspection of the processing metadata reveals that this test '
    'is design-limited. The four groups with stable processing (constant T and hold within group) '
    'all have \u0394d \u2264 8 \u03bcm, too narrow for YS measurement noise to permit a stable slope. '
    'The two groups with meaningful \u0394d (93 \u03bcm and 165 \u03bcm) span recrystallization '
    'temperatures of 700\u20131025 \u00b0C and 825\u20131150 \u00b0C respectively, so their slopes confound '
    'HP with processing-induced microstructural differences. Consequently, no clean within-replicate '
    'test of k_HP composition-dependence is possible from the present dataset; the constant-k '
    'conclusion above rests on the M3 fit quality and the per-alloy k_eff regression, not on '
    'within-composition comparison.'
)

add_heading('7.4 k_HP in Context: Literature Comparison', level=2)
add_para(
    'The fitted k_HP = 766 MPa\u00b7\u03bcm\xb9\u2082 is compared against literature values for FCC '
    'alloys. Our value is 55% higher than the canonical CoCrFeMnNi (Cantor alloy) value of 494 '
    '(Otto et al., 2013) and 2.9\u00d7 higher than equiatomic CoCrNi (265 MPa\u00b7\u03bcm\xb9\u2082; Yoshida et al., 2017). '
    'Notably, Yoshida et al. report that the CoCrNi friction stress is unusually high '
    '(\u03c3\u2080 \u2248 218 MPa, exceeding even CoCrFeMnNi at 125 MPa), so the moderate CoCrNi HP slope coexists '
    'with a large composition-driven intercept. Our elevated k_HP value is physically reasonable given '
    'that our dataset includes V-rich and Al-containing alloys, which are reported to exhibit elevated '
    'k_HP due to ordering effects, reduced stacking fault energy, and enhanced grain-boundary pinning.'
)

add_table(
    ['Alloy', 'k_HP (MPa\u00b7\u03bcm\xb9\u2082)', 'Reference'],
    [
        ['Pure Cu', '110', 'Cordero et al. (2016)'],
        ['Pure Ni', '160', 'Armstrong (2014)'],
        ['316L stainless steel', '322', 'Singh et al. (2002)'],
        ['CoCrFeMnNi', '494', 'Otto et al. (2013)'],
        ['CoCrFeNi', '516', 'Yoshida et al. (2019)'],
        ['CoCrNi', '265', 'Yoshida et al. (2017)'],
        ['This work (93 FCC HEAs)', '766', 'Global constant from M3'],
        ['Al\u2080.\u2083CoCrFeNi (FCC)', '824', 'Gwalani et al. (2017)'],
        ['Al\u2080.\u2083CoCrFeNi (ordered)', '1014', 'Gwalani et al. (2017)'],
    ]
)

add_figure('41_kHP_vs_composition.png',
           'Figure 24. Effective k_HP vs element content for each alloy (colored by grain size). '
           'Linear fits show no significant correlation for any element. Literature comparison in '
           'bottom-right panel.')

add_figure('42_kHP_bayesian_composition.png',
           'Figure 25. Bayesian analysis of k_HP composition dependence: (a) k_HP vs V with 94% '
           'credible interval and literature benchmarks, (b) k_HP vs Mn, (c) posterior distributions '
           'of \u03b2_V and \u03b2_Mn coefficients showing both span zero.')

add_figure('43_kHP_diagnostics.png',
           'Figure 26. k_HP diagnostics: (a) distribution of effective k_HP (red = global constant, '
           'green = CoCrFeMnNi literature), (b) k_eff vs grain size confirming no grain-size dependence, '
           '(c) parity comparison of constant k vs k(V,Mn) models.')
doc.add_page_break()

# ============================================================
# 8. EXHAUSTIVE ML MODEL SEARCH
# ============================================================
add_heading('8. Exhaustive Machine Learning Model Search', level=1)

add_heading('8.1 Models and Methodology', level=2)
add_para(
    'We evaluated 17 models spanning five categories: Linear (OLS, Ridge, ElasticNet), Kernel (SVR, '
    'KRR, GPR), Tree-based (Random Forest, XGBoost, CatBoost, LightGBM), Physics-informed '
    '(M3: composition-dependent Hall\u2013Petch), and Ensemble (Stacking, '
    'Average). The stacking ensemble combines the best model from each of five diverse families '
    '(tree, kernel, linear, compact boosting, physics) via a RidgeCV meta-learner. '
    'Hyperparameter optimization used Optuna with 50 Bayesian trials for all non-trivial '
    'models. Evaluation metrics include LOO R\u00b2, LOBO R\u00b2, AIC, AICc, and BIC.'
)
add_para(
    'Leave-one-out (LOO) cross-validation holds out a single data point and trains on the '
    'remaining 92, repeating for all 93 samples. While standard, LOO can be optimistic when '
    'the dataset contains structured groups. Our 93 alloys were synthesized across six sequential '
    'experimental iterations (BBA, BBB, BBC, CBA, CBB, CBC), each targeting a different region '
    'of the FCC HEA composition space. Data points within the same iteration share systematic '
    'similarities\u2014processing campaign, equipment calibration, grain-size distributions\u2014that '
    'are absent between iterations. Leave-one-batch-out (LOBO) cross-validation removes all '
    'samples from one iteration at a time (6-fold), training on the remaining five and predicting '
    'the held-out batch. This tests whether a model has learned transferable composition\u2013'
    'microstructure\u2013property relationships or has merely memorized batch-specific patterns. '
    'A large LOO\u2013LOBO gap signals overfitting to within-batch correlations; a small or negative '
    'gap (LOBO R\u00b2 \u2265 LOO R\u00b2) indicates robust generalization. For alloy design applications '
    'where the model must predict properties of alloys from future experimental campaigns, '
    'LOBO R\u00b2 is the more honest performance metric.'
)

add_heading('8.2 Complete Results', level=2)

# Read model results
df_models = pd.read_csv(f'{RESULTS_DIR}/model_search_results_v2.csv')
df_models = df_models[df_models['LOO_R2'] > -1].sort_values('LOO_R2', ascending=False)

model_rows = []
for _, r in df_models.iterrows():
    lobo = f"{r['LOBO_R2']:.3f}" if pd.notna(r['LOBO_R2']) else '\u2014'
    bic_val = f"{r['BIC']:.0f}" if pd.notna(r['BIC']) and r['BIC'] < 100000 else f"{r['BIC']:.0f}"
    model_rows.append([
        r['Model'],
        r['Features'],
        str(int(r['n_feat'])),
        f"{r['LOO_R2']:.3f}",
        f"{r['LOO_RMSE']:.1f}",
        lobo,
        str(int(r['k_eff'])),
        bic_val,
    ])

add_table(
    ['Model', 'Features', 'n_f', 'LOO R\u00b2', 'RMSE', 'LOBO R\u00b2', 'k_eff', 'BIC'],
    model_rows
)

add_heading('8.3 Key Findings', level=2)

add_bold_then_text('Best LOO performance: ',
    'XGBoost with INTERACTIONS_ALT features (64 features) achieves LOO R\u00b2 = 0.729, '
    'RMSE = 42.2 MPa, MAE = 27.9 MPa. This represents a 53% reduction in unexplained variance '
    'relative to the simple Hall\u2013Petch baseline (Train R\u00b2 = 0.43), or equivalently a 70% '
    'relative increase in R\u00b2.')

add_bold_then_text('Best parsimony (BIC): ',
    'The stacking ensemble (Ridge meta-learner over 5 base models, including the physics-informed '
    'M3 Hall\u2013Petch model) achieves R\u00b2 = 0.698 with '
    'only k_eff = 6 effective parameters, yielding BIC = 717\u2014the best among all models with '
    'R\u00b2 > 0.5.')

add_bold_then_text('Best generalization (LOBO): ',
    'The stacking ensemble maintains LOBO R\u00b2 = 0.707 > LOO R\u00b2 = 0.698, indicating robust '
    'generalization across batches. However, the stacking LOBO uses LOO (not LOBO) base-model '
    'predictions as meta-features, so its LOBO R\u00b2 may be mildly optimistic. '
    'XGBoost shows a notable LOO\u2013LOBO gap (0.729 vs 0.574), '
    'suggesting partial overfitting to batch-specific patterns.')

add_bold_then_text('Linear models: ',
    'Ridge and ElasticNet achieve LOO R\u00b2 \u2248 0.636\u20130.638 with minimal LOO\u2013LOBO gaps '
    '(\u22480.012), demonstrating that simpler models generalize more reliably but sacrifice peak accuracy.')

add_figure('20_model_comparison_bar.png',
           'Figure 27. LOO R\u00b2 bar chart for all 16 models with BIC annotations.')
add_figure('21_parity_grid.png',
           'Figure 28. Parity plots (predicted vs experimental YS) for all models, colored by batch.')
add_figure('23_loo_vs_lobo.png',
           'Figure 29. LOO vs LOBO R\u00b2 scatter. Points above the diagonal indicate better '
           'generalization to unseen batches than to individual held-out samples.')
doc.add_page_break()

# ============================================================
# 9. SHAP ANALYSIS
# ============================================================
add_heading('9. SHAP Feature Importance Analysis', level=1)

add_heading('9.1 XGBoost SHAP Results', level=2)
add_para(
    'SHAP (SHapley Additive exPlanations) analysis of the best XGBoost model decomposes each '
    'prediction into additive contributions from individual features. The analysis reveals the '
    'following feature importance hierarchy:'
)

add_bold_then_text('\u03a9 (stability parameter): ',
    'The single most influential descriptor. High \u03a9 values are associated with lower yield '
    'strengths, likely because high-\u03a9 alloys tend to have lower lattice distortion and simpler '
    'compositions, resulting in weaker solid-solution strengthening.')

add_bold_then_text('Mn\u00d7d\u207b\xb9 and V\u00d7d\u207b\xb9: ',
    'Composition\u2013grain-size interaction features dominate over their constituent individual '
    'features. This demonstrates that XGBoost leverages composition-dependent Hall\u2013Petch slopes.')

add_bold_then_text('Alternative grain-size features: ',
    'd\u207b\xb9 appears prominently in the top-10 SHAP ranking despite d\u207b\xb9\u2082 also being '
    'available. Tree-based models benefit from multiple grain-size representations because d\u207b\xb9 '
    'compresses the grain-size axis such that fine-grained alloys occupy a wider range, enabling '
    'more informative splits.')

add_bold_then_text('SSS descriptors: ',
    '\u03c3_VLC, \u03b5_Labusch, and \u03b4 contribute moderately, consistent with their role as '
    'composition-dependent friction-stress proxies.')

add_figure('10_shap_summary_YS.png',
           'Figure 30. SHAP beeswarm plot for YS: top 20 features ranked by mean |SHAP|. '
           'Color indicates feature value (red = high, blue = low).')
add_figure('10_shap_bar_YS.png',
           'Figure 31. SHAP mean absolute value bar plot for yield strength.')
add_figure('10_shap_dependence_YS.png',
           'Figure 32. SHAP dependence plots for the six most important features.')

add_heading('9.2 SHAP Interaction Analysis', level=2)
add_para(
    'SHAP interaction values reveal that the strongest pairwise interactions involve grain size and '
    'composition: d\u207b\xb9\u2082\u00d7V, d\u207b\xb9\u2082\u00d7Mn, and d\u207b\xb9\u00d7\u03a9. '
    'These interactions reflect composition-dependent strengthening, but should be interpreted with '
    'care. As shown in Section 7.3, the refined two-stage analysis demonstrates that these '
    'interactions are proxies for composition-dependent \u03c3\u2080, not composition-dependent k_HP. '
    'Additionally, V-containing alloys tend to have finer grains (r(V, d) = \u22120.24, p = 0.019), '
    'so some SHAP interaction values may partially reflect this composition\u2013grain-size correlation '
    'rather than a direct strengthening mechanism.'
)
add_figure('12_shap_interactions_YS.png',
           'Figure 33. SHAP interaction plots for yield strength, showing pairwise feature interactions.')

add_heading('9.3 Hardness SHAP Analysis', level=2)
add_para(
    'The SHAP analysis for hardness (HV) reveals a similar feature hierarchy, with \u03a9 and '
    'grain-size features dominating. A comprehensive analysis of the HV\u2013YS relationship, '
    'including the Tabor relation (C_eff = 5.13 \u00b1 1.36 >> 3), HV Hall\u2013Petch scaling, '
    'and composition-dependent H\u2080 models, is provided in Section 15. Despite the much weaker '
    'Hall\u2013Petch fit for HV (R\u00b2 = 0.14 vs 0.43 for YS), the rank ordering of SHAP '
    'feature importances is preserved, confirming that the same strengthening mechanisms '
    'govern both properties.'
)
add_figure('10_shap_summary_HV.png',
           'Figure 34. SHAP beeswarm plot for Vickers hardness.')
doc.add_page_break()

# ============================================================
# 10. SYMBOLIC REGRESSION
# ============================================================
add_heading('10. Symbolic Regression (PySR)', level=1)

add_para(
    'PySR (Julia-backed symbolic regression) was used to discover interpretable analytical '
    'expressions for yield strength. Three strategies were employed:'
)

add_bold_then_text('Strategy 1 \u2014 Full model: ',
    'Fit \u03c3_y = f(composition, d, processing) directly. Pareto-optimal equations with fewer '
    'than 10 terms capture the dominant strengthening trends.')

add_bold_then_text('Strategy 2 \u2014 Decomposed \u03c3\u2080(comp): ',
    'Fit the friction stress as a function of composition alone, yielding equations of the form '
    '\u03c3\u2080 = g(V_frac, Mn_frac, \u03a9, ...).')

add_bold_then_text('Strategy 3 \u2014 Decomposed k_HP(comp): ',
    'Fit the Hall\u2013Petch coefficient as a function of composition. PySR discovers nonlinear '
    'expressions involving V, Mn, and Cr+Fe terms.')

add_para(
    'The PySR decomposed form \u03c3_y = \u03c3\u2080(comp) + k_HP(comp)\u00b7d\u207b\xb9\u2082 '
    'suggests both \u03c3\u2080 and k_HP may be composition-dependent. PySR\u2019s discovery of '
    'composition-dependent k_HP represents a legitimate alternative decomposition. The constrained '
    'two-stage analysis in Section 7.3 is consistent with composition-independent k_HP '
    '(R\u00b2 = 0.006, F p = 0.999), but the current dataset cannot definitively distinguish between '
    '\u03c3\u2080(comp) + constant k and \u03c3\u2080\u2032(comp) + k(comp)\u00b7d\u207b\xb9\u2082, '
    'as both achieve similar overall fit. PySR optimizes global fit without constraining the '
    '\u03c3\u2080/k decomposition, which may allow composition effects to redistribute between terms. '
    'Single-composition grain-size series would be needed to resolve this ambiguity definitively.')

add_figure('15_pysr_results.png',
           'Figure 35. PySR results: (a) full model predictions, (b) \u03c3\u2080(comp) fit, '
           '(c) k_HP(comp) fit, (d) combined model parity plot.')
add_figure('16_pysr_pareto.png',
           'Figure 36. PySR Pareto front: equation complexity vs prediction loss. Diminishing returns '
           'beyond complexity ~15.')
doc.add_page_break()

# ============================================================
# 11. SISSO SYMBOLIC REGRESSION AND JIANG MODEL COMPARISON
# ============================================================
add_heading('11. SISSO Symbolic Regression and Jiang Model Comparison', level=1)

add_para(
    'SISSO (Sure Independence Screening and Sparsifying Operator) is a compressed-sensing approach '
    'to symbolic regression that constructs an expanded feature space via algebraic combinations of '
    'input descriptors, then selects the best low-dimensional (additive) descriptor via \u2113\u2080 '
    'regularization (Ouyang et al., Phys. Rev. Materials, 2018). Unlike PySR, which optimizes '
    'expression trees via genetic programming, SISSO guarantees that the selected expression is '
    'optimal within the constructed feature space for a given dimensionality.'
)

add_heading('11.1 Physics-Informed Feature Space', level=2)
add_para(
    'We use 27 Oliynyk-style descriptors (Oliynyk et al., Chem. Mater., 2016) computed from '
    'elemental properties: for each of atomic radii, shear modulus, electronegativity, melting '
    'point, and bulk modulus, we compute the composition-weighted mean, variance, delta (\u03c3/\u03bc), '
    'and range. These are supplemented by VEC, HEA thermodynamic descriptors (\u03b4, \u0394S_mix, '
    '\u0394H_mix, \u03a9), physics-based SSS estimates (\u03a6_VLC, \u03b5_Labusch, \u03c3_TC), and '
    'd\u207b\xb9\u2082 for Hall\u2013Petch scaling\u201428 features total. SISSO expands these via '
    'one tier of algebraic operators (+, \u2212, \u00d7, \u00f7), generating ~15,000 candidate features, '
    'then selects the best 3-dimensional (3-term additive) descriptor via \u2113\u2080 regularization.'
)

add_heading('11.2 SISSO Discovered Equation', level=2)
add_para(
    'The best SISSO equation (4 effective parameters: 3 coefficients + intercept) is:'
)
add_para(
    '\u03c3_y = 120.5\u00b7(\u03c3\u00b2_r / r_range) + 9356\u00b7(d\u207b\xb9\u2082 / \u0394S_mix) '
    '+ 1134\u00b7(\u03c3\u00b2_\u03c7 / \u03b4_\u03bc) \u2212 43.3',
    bold=True
)
add_para(
    'where \u03c3\u00b2_r and r_range are the composition-weighted variance and range of atomic radii, '
    '\u03c3\u00b2_\u03c7 is the electronegativity variance, and \u03b4_\u03bc is the shear modulus mismatch '
    'parameter. This equation achieves LOO R\u00b2 = 0.665 (RMSE = 46.9 MPa), LOBO R\u00b2 = 0.380, '
    'and BIC = 714\u2014the lowest BIC among all 23 models evaluated.'
)
add_para(
    'Each term carries physical meaning: (1) the first captures solid-solution strengthening through '
    'atomic size mismatch variance normalized by the total radius span; (2) the second encodes '
    'Hall\u2013Petch grain-boundary strengthening modulated by configurational entropy, implying that '
    'alloys with higher mixing entropy require less grain-boundary strengthening per unit d\u207b\xb9\u2082; '
    'and (3) the third reflects electronic contributions through electronegativity variance '
    'normalized by shear modulus mismatch.'
)

add_heading('11.3 Three Modeling Strategies', level=2)
add_para(
    'Three SISSO strategies were evaluated, mirroring the PySR decomposition:'
)
add_bold_then_text('Strategy 1 \u2014 Full model: ',
    'SISSO predicts \u03c3_y directly from all 28 features. This yields the best result (LOO R\u00b2 = 0.665, '
    'BIC = 714).')
add_bold_then_text('Strategy 2 \u2014 Decomposed \u03c3\u2080(comp): ',
    'SISSO learns \u03c3\u2080 from M3 residuals; final = SISSO_\u03c3\u2080 + k_HP\u00b7d\u207b\xb9\u2082. '
    'LOO R\u00b2 = 0.595.')
add_bold_then_text('Strategy 3 \u2014 Decomposed k_HP(comp): ',
    'SISSO learns k_HP(composition); final = \u03c3\u2080_mean + SISSO_k_HP\u00b7d\u207b\xb9\u2082. '
    'LOO R\u00b2 = 0.614.')

add_figure('57_sisso_full_parity.png',
           'Figure 37. SISSO full model: LOO predictions vs observed yield strength, colored by batch. '
           'The 4-parameter equation achieves LOO R\u00b2 = 0.665 (RMSE = 46.9 MPa).')
add_figure('58_sisso_sigma0_parity.png',
           'Figure 38. SISSO decomposed \u03c3\u2080 model: LOO R\u00b2 = 0.595.')
add_figure('59_sisso_khp_parity.png',
           'Figure 39. SISSO decomposed k_HP model: LOO R\u00b2 = 0.614.')
add_figure('60_sisso_complexity.png',
           'Figure 40. Complexity\u2013accuracy tradeoff: SISSO vs PySR vs ML models. '
           'SISSO achieves the best BIC with only 4 parameters.')

add_heading('11.4 Jiang et al. Pure-Metal Hall\u2013Petch Model', level=2)
add_para(
    'We test the transferability of the pure-metal Hall\u2013Petch model of Jiang et al. (2022), '
    'which uses cohesive energy (W), thermal expansion (\u03b1), grain boundary energy (\u03b3_GB), '
    'and a valence electron density parameter (S_VED) to predict both \u03c3\u2080 and k_HP for '
    'elemental metals: \u03c3_y = 79\u00b7W/(S\u00b3\u221a\u03b1) + 1.2\u00b7\u221a(\u03b3\u00b7E/\u03b1)\u00b7d\u207b\xb9\u2082.'
)
add_para(
    'Applied directly with rule-of-mixtures-averaged elemental properties, the Jiang model yields '
    'R\u00b2 = \u22122.23 (fails completely). Even after recalibration (fitting coefficients a, b, and '
    'an intercept to our dataset), the LOO R\u00b2 reaches only 0.346\u2014worse than the baseline '
    'Hall\u2013Petch (R\u00b2 = 0.406). The failure stems from the Jiang descriptors being '
    'rule-of-mixtures means of elemental properties, which are nearly identical across our 93 '
    'Ni-rich FCC alloys (\u0394W \u00b112%, \u0394\u03b1 \u00b18%, \u0394\u03b3_GB \u00b19%).'
)
add_para(
    'The contrast between Jiang (R\u00b2 = 0.346) and SISSO (R\u00b2 = 0.665) illustrates a '
    'fundamental principle for HEA modeling: variance-based descriptors (\u03c3\u00b2_r varies by 30\u00d7, '
    '\u03c3\u00b2_\u03c7 by 15\u00d7 across the dataset) provide far greater discriminating power than '
    'rule-of-mixtures means in concentrated solid solutions, where all alloys share similar average '
    'properties by design.'
)

add_figure('61_jiang_comparison.png',
           'Figure 41. Comparison of (a) Jiang et al. pure-metal Hall\u2013Petch applied directly '
           '(R\u00b2 = \u22122.23), (b) recalibrated Jiang model (LOO R\u00b2 = 0.346), and (c) SISSO '
           'full model (LOO R\u00b2 = 0.665). The Jiang model\u2019s rule-of-mixtures descriptors fail to '
           'differentiate HEA compositions.')

add_heading('11.5 Expanded SISSO Search (v2) and EML Symbolic Regression', level=2)
add_para(
    'To verify that the SISSO equation (Eq. above) is robust to search-space choices, we ran an '
    'expanded search (SISSO v2) with three enhancements: (1) the fixed d\u207b\xb9\u2082 grain-size '
    'feature is replaced by 10 flexible alternatives (d\u207b\u1d45 for \u03b1 \u2208 {1/3, 0.40, '
    '0.45, 0.50, 0.55, 0.60, 2/3, 0.75, 1.0} plus ln(d)); (2) unary operators (x\u00b2, \u221ax, '
    'x\u207b\xb9) are added alongside the binary operators; (3) the dimension D \u2208 {1,...,4} is '
    'selected by BIC rather than fixed at 3. The SIS threshold is increased from 20 to 30, '
    'yielding 36 primary features and 3,293 after one tier of expansion.'
)
add_para(
    'Despite the substantially larger search space, SISSO v2 achieves lower LOO R\u00b2 = 0.604 '
    '(RMSE = 51.0 MPa) and lower LOBO R\u00b2 = 0.350 compared to the constrained v1 search '
    '(LOO R\u00b2 = 0.665, LOBO R\u00b2 = 0.380). BIC-based dimension selection chooses dim = 4 '
    '(BIC = 714.7), though dim = 3 is near-equivalent (\u0394BIC = 2.1). LOO was evaluated at '
    'dim = 3 due to computational cost (dim = 4 SO step takes ~580 s/fold). The full-data '
    'equation at dim = 3 is: \u03c3_y = 15071\u00b7(EN_var/\u03a9) + 53.7\u00b7(VEC_mean \u2212 ln d) '
    '+ 38250\u00b7(\u03b4_r/r_range) \u2212 142.2.'
)
add_para(
    'The expanded search reveals two insights: (1) at dimension 1, the best single feature is '
    'd\u207b\u2070\u00b7\u2074\u2075/\u03a9 rather than d\u207b\xb9\u2082, and higher-dimensional equations '
    'prefer ln(d) over d\u207b\xb9\u2082, confirming that the optimal grain-size exponent is not sharply '
    'defined (consistent with the BIC-indistinguishable scaling laws of Section 5); (2) however, '
    'different LOO folds select different grain-size features and descriptor combinations, '
    'introducing feature-selection instability that degrades cross-validated performance. '
    'Constraining the search to d\u207b\xb9\u2082 alone stabilizes model selection across folds '
    'and yields superior generalization\u2014a concrete example where inductive bias outweighs '
    'search flexibility in small-sample symbolic regression.'
)
add_para(
    'We also evaluated EML (Elementary function as a single operation; Odrzywo\u0142ek, 2025), '
    'a universal-operator symbolic regression approach where a single binary operator '
    'eml(x,y) = exp(x) \u2212 ln(y) can represent all elementary functions when composed into '
    'a binary tree. At tree depth 1 (2 leaves, 20 effective parameters), EML achieves LOO '
    'R\u00b2 = 0.321 (BIC = 772); deeper trees (depth 2\u20133, 38\u201374 parameters) overfit '
    'severely (LOO R\u00b2 < 0). The non-convex loss landscape with many local minima makes '
    'EML unsuitable for datasets of this size (n = 93). Both experiments confirm that the '
    'constrained SISSO search with d\u207b\xb9\u2082 and binary operators is near-optimal.'
)
add_figure('63b_sisso_v2_parity.png',
           'Figure 42. SISSO v2 expanded search: LOO predictions vs observed yield strength. '
           'Despite a larger feature space, LOO R\u00b2 = 0.604 is lower than the constrained '
           'v1 search (R\u00b2 = 0.665) due to feature-selection instability.')
add_figure('63a_sisso_v2_bic_landscape.png',
           'Figure 43. BIC landscape across SISSO v2 dimensions 1\u20134. BIC decreases monotonically '
           'with diminishing returns; dimensions 3 and 4 are BIC-equivalent (\u0394BIC = 1.7).')

add_heading('11.6 Robustness Analysis: Singularity in the SISSO Equation', level=2)
add_para(
    'The third term of the SISSO equation\u2014\u03c3\u00b2_\u03c7/\u03b4_\u03bc\u2014contains a latent '
    'singularity: when alloying elements have similar shear moduli (\u03bc\u1d62 \u2248 \u03bc\u0304), '
    'the denominator \u03b4_\u03bc \u2192 0 and predictions diverge. Among the eight elements in our system, '
    'Co, Fe, Mn, and Ni have \u03bc \u2208 [75, 82] GPa, so any equiatomic ternary from this group '
    'has \u03b4_\u03bc < 0.05, producing predictions exceeding 1000 MPa.'
)
add_para(
    'To obtain a robust alternative, we re-ran SISSO after removing \u03b4_\u03bc from the candidate '
    'feature pool. Three independent searches\u2014varying which delta-type features are excluded\u2014'
    'all converge to the same equation:'
)
add_para(
    '\u03c3_y = 113.0\u00b7(\u03c3\u00b2_r / r_range) + 9837\u00b7(d\u207b\xb9\u2082 / \u0394S_mix) '
    '+ 5437\u00b7(\u03c3\u00b2_\u03c7 \u2212 \u03a6_VLC) \u2212 27.0',
    bold=True
)
add_para(
    'This robust equation achieves LOO R\u00b2 = 0.609 (RMSE = 50.7 MPa) and BIC = 717, with '
    'physically reasonable predictions (168\u2013614 MPa) for all 56 equiatomic ternaries\u2014'
    'compared to [169, 1524] MPa for the original equation. The first two terms are nearly '
    'identical to the original; the third replaces the singular ratio \u03c3\u00b2_\u03c7/\u03b4_\u03bc '
    'with the bounded difference \u03c3\u00b2_\u03c7 \u2212 \u03a6_VLC, coupling electronegativity variance '
    'with VLC lattice mismatch. The modest accuracy cost (\u0394R\u00b2_LOO = 0.045) is justified by '
    'the elimination of catastrophic extrapolation failures (see External Validation section).'
)
doc.add_page_break()

# ============================================================
# 12. INFORMATION CRITERIA
# ============================================================
add_heading('12. Information Criteria and Model Selection', level=1)

add_heading('12.1 AIC and BIC Framework', level=2)
add_para(
    'Model comparison employs the Akaike Information Criterion (AIC = n\u00b7ln(RSS/n) + 2k) and '
    'Bayesian Information Criterion (BIC = n\u00b7ln(RSS/n) + k\u00b7ln(n)), where n = 93 and k is '
    'the effective parameter count. BIC penalizes complexity more heavily than AIC at this sample '
    'size (ln(93) = 4.53 vs 2). For tree-based models, k_eff is estimated as the total number of '
    'leaves across all trees. BIC comparisons across model families are approximate because k_eff '
    'estimation methods differ (regression parameters for linear models vs total leaves for tree '
    'ensembles), but the qualitative ranking is robust.'
)

add_heading('12.2 BIC Rankings', level=2)
add_table(
    ['Rank', 'Model', 'BIC', 'k_eff', 'LOO R\u00b2'],
    [
        ['1', 'SISSO Full (Eq.)', '714', '4', '0.665'],
        ['1b', 'SISSO Robust', '717', '4', '0.609'],
        ['2', 'SISSO v2 (expanded)', '717', '4', '0.604'],
        ['3', 'Stacking (Ridge)', '717', '6', '0.698'],
        ['4', 'Average Ensemble', '728', '5', '0.700'],
        ['5', 'M3 (\u03c3\u2080 all elem)', '743', '10', '0.652'],
        ['6', 'EML depth-1', '772', '20', '0.321'],
        ['7', 'Jiang (recal.)', '785', '3', '0.346'],
        ['8', 'ElasticNet', '839', '32', '0.637'],
        ['9', 'Ridge', '847', '34', '0.638'],
        ['10', 'SVR (RBF)', '1,062', '83', '0.654'],
        ['11', 'XGBoost', '16,889', '3,749', '0.729'],
    ]
)

add_para(
    'SISSO achieves the best BIC (714) with only 4 parameters, followed closely by the stacking '
    'ensemble (BIC = 717, 6 parameters) and the expanded SISSO v2 search (BIC = 717, 4 parameters). '
    'Notably, SISSO v2 has BIC nearly identical to v1 despite a larger feature space, but its '
    'lower LOO R\u00b2 (0.604 vs 0.665) reflects feature-selection instability across folds. '
    'EML symbolic regression (BIC = 772, 20 parameters) and Jiang et al. (BIC = 785) both rank '
    'below the physics-informed SISSO equation. XGBoost, despite its highest LOO R\u00b2, has '
    'k_eff = 3,749, yielding the worst BIC among well-performing models.'
)

add_figure('26_model_ic_comparison.png',
           'Figure 44. Information criteria comparison across models: AIC, BIC, and LOO R\u00b2.')

add_heading('12.3 Model Selection as a Design Choice', level=2)

add_para(
    'The preceding BIC analysis, combined with LOO/LOBO accuracy (Section 8) and external '
    'validation (Section 13), reveals that no single model simultaneously optimizes all evaluation '
    'criteria. Rankings shift depending on the protocol, reflecting genuinely different aspects of '
    'model quality.'
)

# Unified comparison table
add_table(
    ['Model', 'LOO R\u00b2', 'LOBO R\u00b2', 'BIC', 'Ext. RMSE (MPa)'],
    [
        ['XGBoost', '0.729', '0.574', '16,889', '\u2014'],
        ['SISSO Full', '0.665', '\u2014', '714', '421'],
        ['SISSO Robust', '0.609', '\u2014', '717', '122'],
        ['M3 (comp-HP)', '0.652', '0.625', '743', '133'],
        ['Stacking', '0.698', '0.707', '717', '\u2014'],
    ],
)

add_para(
    'XGBoost achieves the highest LOO R\u00b2 but suffers the largest LOO\u2013LOBO degradation '
    '(\u0394R\u00b2 = 0.155), partly attributable to hyperparameter optimization on the full dataset. '
    'SISSO Full attains the best BIC with only 4 parameters, yet its \u03c3\u00b2_\u03c7/\u03b4_\u03bc singularity '
    'renders it unsuitable for deployment on compositions with near-zero shear-modulus mismatch. '
    'The robust variant sacrifices 0.045 in LOO R\u00b2 but reduces external RMSE from 421 to '
    '122 MPa\u2014the best generalization among all models tested.'
)

add_para(
    'In practice, this multiplicity reflects genuinely different priorities in alloy design:'
)

selection_guidance = [
    ('Peak in-sample accuracy: ', 'XGBoost (caveat: HPO bias inflates LOO; largest LOBO drop)'),
    ('Closed-form interpretability + best external generalization: ',
     'SISSO Robust (4 parameters, RMSE = 163 MPa on 82 external points)'),
    ('Balanced accuracy\u2013generalization: ',
     'Stacking ensemble (highest LOBO R\u00b2 = 0.707; competitive LOO and BIC)'),
    ('Physical decomposition (\u03c3\u2080 vs k_HP): ',
     'M3 (direct window into friction-stress versus grain-boundary partition; RMSE = 133 MPa externally)'),
    ('Best parsimony (lowest BIC): ',
     'SISSO Full (but deployment-unsafe due to singularity\u2014use Robust variant instead)'),
]
for bold_part, normal_part in selection_guidance:
    add_bold_then_text(bold_part, normal_part)

add_para(
    'Reporting all four metrics\u2014rather than a single leaderboard ranking\u2014enables practitioners '
    'to select the model best suited to their specific deployment context.'
)

doc.add_page_break()

# ============================================================
# 13. EXTERNAL VALIDATION
# ============================================================
add_heading('13. External Validation Against Independent Literature Data', level=1)

add_para(
    'To test out-of-sample generalization, SISSO Full, SISSO Robust, and M3 are evaluated on '
    '82 independent data points from four literature sources: (1) the Citrine/Borg MPEA dataset '
    '(48 entries, aggregated, heterogeneous quality); (2) Schneider et al. CrFeNi compression '
    'data (6 entries, verified from open-access paper); (3) Otto et al. CoCrFeMnNi tension '
    'data (3 entries, HP-derived from published parameters); and (4) Huang et al. Vickers hardness '
    'data converted to yield strength (25 entries, C_eff = 5.13).'
)

add_heading('13.1 Overall Performance', level=2)
add_table(
    ['Model', 'R\u00b2', 'RMSE (MPa)', 'MAE (MPa)', 'Bias (MPa)'],
    [
        ['SISSO Full', '\u221214.8', '421', '220', '+144'],
        ['SISSO Robust', '\u22120.33', '122', '82', '+5'],
        ['M3', '\u22120.58', '133', '97', '\u221264'],
    ]
)
add_para(
    'SISSO Full exhibits catastrophic failure (RMSE = 421 MPa), driven entirely by the '
    '\u03c3\u00b2_\u03c7/\u03b4_\u03bc singularity: for Mn-containing alloys (CoNiMn, CoFeNiMn) where '
    '\u03b4_\u03bc < 0.04, predicted \u03c3\u2080 exceeds 600 MPa. SISSO Robust eliminates this pathology, '
    'achieving RMSE = 163 MPa\u2014outperforming M3 (RMSE = 133 MPa, '
    'bias = \u221264 MPa) despite having only 4 parameters vs 10.'
)

add_heading('13.2 Parity Plots', level=2)
add_figure('67_external_parity.png',
           'Figure 45. External validation parity plots: predicted vs measured yield strength for '
           'SISSO Full (left), SISSO Robust (center), and M3 (right). SISSO Full catastrophically '
           'overpredicts for compositions with low \u03b4_\u03bc. SISSO Robust achieves the best overall '
           'performance (RMSE = 163 MPa).')

add_heading('13.3 Hall\u2013Petch Slope Comparison', level=2)
add_para(
    'For alloys with \u22653 grain sizes, we fit YS = \u03c3\u2080 + k_HP\u00b7d\u207b\xb9\u2082 '
    'to compare experimental and predicted Hall\u2013Petch slopes.'
)
add_figure('68_external_hp_slopes.png',
           'Figure 46. Hall\u2013Petch slopes for 12 alloy compositions. Experimental data (black), '
           'SISSO Full (blue), SISSO Robust (green), M3 (orange). SISSO Robust closely tracks M3 '
           'with k_HP \u2248 735\u20131077 MPa\u00b7\u03bcm\xb9\u2082.')

add_heading('13.4 Error Metrics by Source', level=2)
add_figure('69_external_error_bars.png',
           'Figure 47. Error metrics (R\u00b2, RMSE, MAE) by data source and model. '
           'SISSO Robust (green) consistently outperforms SISSO Full (blue) and performs '
           'comparably to M3 (orange).')

add_heading('13.5 Residual Distributions', level=2)
add_figure('70_external_residuals.png',
           'Figure 48. Residual distributions (Exp \u2212 Pred) for each model by data source. '
           'SISSO Full shows heavy positive tails from the singularity; SISSO Robust and M3 '
           'are approximately centered around zero.')

add_para(
    'Data quality caveats: The Schneider data is from compression tests (not tension), '
    'the Otto data is derived from published HP parameters rather than raw data, '
    'and the Huang HV\u2192YS conversion uses our estimated C_eff = 5.13. '
    'These heterogeneities limit quantitative interpretation but the qualitative conclusion\u2014'
    'that SISSO Robust eliminates the singularity pathology while maintaining competitive '
    'accuracy\u2014is robust to these caveats.'
)
doc.add_page_break()

# ============================================================
# 14. LIMITATIONS AND CAVEATS
# ============================================================
add_heading('14. Limitations and Caveats', level=1)

add_para(
    'The following limitations and caveats should be considered when interpreting the results of '
    'this analysis. A series of robustness diagnostics (VIF, Monte Carlo sensitivity, subset analysis, '
    'bootstrap confidence intervals, and Simpson\'s paradox checks) were performed to quantify the '
    'impact of these limitations.'
)

add_heading('14.1 Data Limitations', level=2)

add_para(
    'The dataset comprises N = 93 FCC HEAs spanning 8 elements, yielding a relatively sparse coverage '
    'of the 7-dimensional composition space (after removing Ni as the balance element). Nine compositions '
    'have 2\u20134 data points at different grain sizes; the remaining 72 compositions have a single '
    'measurement each, limiting the ability to estimate per-alloy Hall\u2013Petch parameters.'
)

add_para(
    'Grain sizes are reported as the mean of an EBSD-measured distribution with within-alloy '
    'coefficient of variation CV \u2248 0.74. While this represents the inherent grain size distribution '
    'rather than measurement error of the mean, the uncertainty in the reported mean grain size '
    'propagates into all models. Monte Carlo sensitivity analysis (1,000 replicates) perturbs each '
    'alloy\'s grain size by SD_GS/\u221aN_grains, conservatively assuming N_grains = 10 per EBSD scan '
    '(if 50+ grains were measured, the perturbation and resulting shifts would be ~2\u00d7 smaller). '
    'Under this assumption, M3 coefficients are moderately sensitive to grain-size '
    'uncertainty: the Al coefficient \u03b1_Al shifts by up to 159%, and the Hall\u2013Petch coefficient '
    'k shifts by ~39%. Other element coefficients (Co, Cr, Fe) shift by less than 15%.'
)

add_figure('44_mc_grain_size_sensitivity.png',
           'Figure 49. Monte Carlo grain-size sensitivity analysis: (a) violin plots of \u03c3\u2080 '
           'composition coefficients across 1,000 MC replicates with grain-size perturbation, '
           '(b) distribution of the Hall\u2013Petch coefficient k.')

add_para(
    'A potential selection bias exists: V-containing alloys tend to have finer grains '
    '(r(V_frac, d) = \u22120.242, p = 0.019), likely reflecting alloy design choices rather than a '
    'causal grain-refinement effect of V. This correlation means that V\'s apparent strengthening '
    'effect may partly reflect the correlation between V content and grain refinement in the dataset '
    '(Simpson\'s paradox). The partial correlation of V with YS controlling for d\u207b\xb9\u2082 is '
    'r = +0.589 (vs raw r = +0.641), indicating ~8% attenuation\u2014modest but non-negligible. '
    'The V coefficient \u03b1_V changes by +73% when d\u207b\xb9\u2082 is removed from the model, '
    'confirming some confounding between V content and grain size. Note that the partial correlation '
    'approach captures only linear confounding; nonlinear confounding channels remain unquantified.'
)

add_heading('14.2 Model Limitations', level=2)

add_para(
    'The composition-dependent \u03c3\u2080 model (M3) assumes additive (linear) element contributions '
    'to the friction stress, with no interaction terms. The design matrix condition number is 87.9 '
    '(moderate), and all element variance inflation factors (VIF) are below 5.3, confirming acceptable '
    'multicollinearity among composition features. The d\u207b\xb9\u2082 feature has VIF = 16.8 within '
    'the 64-feature INTERACTIONS_ALT pool, but this inflation is largely artifactual: that pool '
    'contains element\u00d7d\u207b\xb9\u2082 cross-terms together with d\u207b\xb9, d\u207b\xb9\u02e3\u2083, '
    'd\u207b\xb2\u02e3\u2083, and ln(d)/d, all of which are mathematical functions of d\u207b\xb9\u2082. '
    'Against the physically independent feature set\u2014the eight elemental fractions plus three '
    'processing variables\u2014the variance is much milder: R\u00b2(d\u207b\xb9\u2082 ~ comp) = 0.40 '
    'and R\u00b2(d\u207b\xb9\u2082 ~ comp + processing) = 0.54, implying VIF \u2248 1.7\u20132.2 and '
    'leaving ~46% of d\u207b\xb9\u2082 variance independent of composition and processing. The '
    '\u03c3\u2080/k_HP partition is therefore reasonably identifiable in this dataset. Independently, '
    'the k_HP estimate remains sensitive to grain-size measurement uncertainty (Monte Carlo analysis '
    'shows \u00b139% shifts under GS perturbation), and bootstrap confidence intervals reveal that the '
    'intercept and k_HP have SE ratios of ~1.5 (bootstrap SE / OLS SE), indicating departures from '
    'OLS normality assumptions for these parameters.'
)

add_para(
    'The claim that k_HP is not composition-dependent is based on N = 93 alloys and may lack '
    'statistical power to detect weak effects. Subset analysis shows variation in k_HP across '
    'composition groups (V-containing: k \u2248 1269, V-free: k \u2248 299), though bootstrap '
    'confidence intervals overlap. Importantly, these subset fits use simple Hall\u2013Petch '
    '(\u03c3\u2080 + k\u00b7d\u207b\xb9\u2082) without composition-dependent \u03c3\u2080, so the higher '
    'k in V-containing alloys partly reflects V\'s large positive \u03b1_V being absorbed into the '
    'slope. Per-alloy k_HP estimates for the 9 compositions with multiple grain sizes are unreliable '
    'for most cases because 7 of 9 compositions have GS CV < 15% (data points at essentially '
    'identical grain sizes), producing numerically unstable slope estimates. The two compositions with '
    'adequate grain-size variation yield k \u2248 731 and k \u2248 2,302, but these two groups '
    'are not honest local k_HP estimates: both span recrystallization temperatures of ~325 \u00b0C '
    'within the group (700\u20131025 \u00b0C for the k \u2248 2,302 group; 825\u20131150 \u00b0C for '
    'the k \u2248 731 group), so their slopes confound Hall\u2013Petch with processing-induced '
    'microstructural differences (see Section 7.3 for the full within-replicate diagnosis). The 3:1 '
    'spread therefore cannot be interpreted as evidence for or against composition-dependent k_HP. '
    'The constant-k conclusion depends on the M3 decomposition being correct, and is supported by '
    'the aggregate per-alloy k_eff regression (R\u00b2 = 0.006), not by within-composition comparison.'
)

add_figure('45_subset_kHP.png',
           'Figure 50. Subset k_HP consistency: Hall\u2013Petch coefficient estimated within composition '
           'subsets, with bootstrap 95% confidence intervals. Red dashed line = global k = 766.')

add_figure('46_per_alloy_kHP.png',
           'Figure 51. Per-alloy k_HP for the 9 compositions with multiple grain-size measurements. '
           'Green = reliable (sufficient GS variation); gray = low GS variation.')

add_heading('14.3 Interpretation Caveats', level=2)

add_para(
    'The \u03c3\u2080 composition coefficients (\u03b1\u1d62) are empirical regression weights, not '
    'physics-based solid-solution strengthening parameters. They capture the net effect of each element '
    'on strength at infinite grain size (d \u2192 \u221e), which includes solid-solution strengthening, '
    'short-range ordering, stacking-fault energy modifications, and any other composition-dependent '
    'effects. The negative coefficients for Fe (\u2212360 MPa/fraction) and Cu (\u2212334 MPa/fraction) '
    'should be interpreted relative to the Ni-solvent reference, not as absolute weakening effects.'
)

add_para(
    'Element effects may be confounded with processing-induced microstructural differences. The '
    'experimental design partially confounds composition with batch (processing conditions), though '
    'the LOBO evaluation provides some assessment of cross-batch generalizability. The '
    'composition\u2013grain-size correlations noted above (particularly for V) mean that some apparent '
    'composition effects could be artifacts of systematic grain-size differences across alloy families. '
    'Furthermore, all data originate from a single research group with consistent processing protocols; '
    'inter-laboratory validation would strengthen the generalizability of these conclusions.'
)

add_heading('14.4 Robustness Summary', level=2)

add_table(
    ['Diagnostic', 'Result', 'Assessment'],
    [
        ['VIF (elements)', 'All < 5.3', 'Acceptable'],
        ['VIF (d\u207b\xb9\u2082) in 64-feat. pool', '16.8', 'Artifactual (collinear grain-size transforms)'],
        ['VIF (d\u207b\xb9\u2082) vs. comp + proc', '~2.2', 'Honest measure; \u03c3\u2080/k_HP is well-identified'],
        ['Condition number', '87.9', 'Moderate; acceptable'],
        ['MC sensitivity (k)', '\u00b139% shift', 'Moderately sensitive to GS uncertainty'],
        ['MC sensitivity (elements)', '<15% (Co,Cr,Fe)', 'Robust for most elements'],
        ['Bootstrap SE ratios', '0.99\u20131.54', 'Some departure from normality'],
        ['Subset k_HP range', '299\u20131,388', 'Wide range but overlapping CIs'],
        ['Simpson\'s paradox (V)', '8% attenuation', 'Modest confounding'],
        ['V coeff. sensitivity', '+73% without d\u207b\xb9\u2082', 'V confounded with grain size'],
    ]
)

add_figure('47_bootstrap_ci.png',
           'Figure 52. Bootstrap confidence intervals for M3 coefficients: (a) element composition '
           'coefficients with OLS 95% CI (light blue) and bootstrap 95% CI (dark blue), '
           '(b) intercept and k_HP.')

doc.add_page_break()

# ============================================================
# 15. HARDNESS ANALYSIS
# ============================================================
add_heading('15. Hardness Analysis', level=1)

add_para(
    'While the preceding sections focused on yield strength (YS), the dataset also contains '
    'Vickers hardness (HV) for all 94 alloys. This section provides a comprehensive analysis of '
    'the HV\u2013YS relationship (Tabor relation), Hall\u2013Petch scaling laws for HV, and '
    'composition-dependent H\u2080 models, complementing the YS-focused results.'
)

add_heading('15.1 Tabor Framework: HV\u2013YS Coupling', level=2)
add_para(
    'Tabor\u2019s classical result H_V \u2248 3\u00b7\u03c3_y applies to a rigid\u2013perfectly-plastic medium under '
    'sharp indentation; the factor 3 is the plastic constraint factor that slip-line analysis '
    'assigns to the deformation zone beneath the indenter. Real metals strain-harden, so the '
    'correct generalization is H_V \u2248 3\u00b7\u03c3_f(\u03b5_r), where \u03c3_f(\u03b5) is the flow stress at plastic '
    'strain \u03b5 and \u03b5_r is the representative indentation strain. For Vickers geometry, Tabor '
    'estimated \u03b5_r \u2248 0.08; FEM analyses of sharp indentation (Dao et al. 2001) place \u03b5_r in '
    'the range 0.07\u20130.10, and we adopt \u03b5_r = 0.08 throughout.'
)
add_para(
    'Defining the effective Tabor factor C_eff = H_V/\u03c3_y, this gives C_eff = 3\u00b7\u03c3_f(\u03b5_r)/\u03c3_y, '
    'so C_eff is a per-alloy measurement of the flow-stress ratio. For a power-law hardening '
    'material with \u03c3 = K\u00b7\u03b5\u207f and \u03c3_y at the 0.2% offset, \u03c3_f(\u03b5_r)/\u03c3_y = (\u03b5_r/0.002)\u207f = 40\u207f, '
    'yielding C_eff = 3\u00b740\u207f, which inverts to an effective hardening exponent '
    'n_eff = ln(C_eff/3)/ln(40). Because Vickers indentation samples only the early post-yield '
    'regime (0.002 < \u03b5 < 0.08), n_eff is the early-strain hardening exponent and need not '
    'equal the asymptotic Hollomon n fitted across the full stress\u2013strain curve.'
)
add_para(
    'Results: C_eff = 5.13 \u00b1 1.36 (mean \u00b1 std), significantly exceeding the classical value of 3 '
    '(t = 15.0, p < 10\u207b\u00b2\u2077). The Tabor inversion yields n_eff = ln(5.13/3)/ln(40) \u2248 0.15, '
    'corresponding to \u03c3_f(0.08)/\u03c3_y \u2248 1.71 averaged across the dataset. The 1\u03c3 spread of C_eff '
    'propagates to n_eff = 0.15 \u00b1 0.07, spanning hardening behavior from near rigid\u2013perfectly-'
    'plastic (n_eff \u2248 0.08) to strongly hardening (n_eff \u2248 0.22). This is lower than full-range '
    'Hollomon fits reported for representative FCC HEAs at room temperature (typically n = 0.3\u20130.5), '
    'consistent with Vickers sampling the early post-yield regime where hardening rate is below '
    'the asymptotic value. The composition dependence of C_eff is modest in aggregate (R\u00b2 = 0.357), '
    'with V (r = \u22120.47) the strongest single correlate: V raises \u03c3_y via \u03c3_0 (M3 coefficient '
    '\u03b1_V = +291 MPa) but does not proportionally raise \u03c3_f(0.08), so V-rich alloys appear stronger '
    'than they appear hard. C_eff also depends weakly on grain size (r = \u22120.39 vs d\u207b\xb9\u00b2): '
    'this is the shared Hall\u2013Petch correlation with both HV and YS, not a true sensitivity '
    'difference, as the rank-analysis in Section 15.5 makes explicit.'
)
add_figure('50_tabor_relation.png',
           'Figure 53. Tabor relation analysis. (a) HV(MPa) vs YS with C=3 line and best fit. '
           '(b) Distribution of C_eff. (c) C_eff vs grain size. (d) C_eff vs strongest correlated element.')
add_figure('51_tabor_composition.png',
           'Figure 54. Composition dependence of the effective Tabor factor for all 8 elements.')

add_heading('15.2 Hall\u2013Petch for Hardness', level=2)
add_para(
    'The baseline Hall\u2013Petch fit for HV yields H\u2080 = 86.7 HV, k_H = 306.2 HV\u00b7\u03bcm\xb9\u2082 '
    '(= 3003 MPa\u00b7\u03bcm\xb9\u2082), with Train R\u00b2 = 0.169 and LOO R\u00b2 = 0.136. This is substantially '
    'weaker than the YS Hall\u2013Petch fit (R\u00b2 = 0.43), indicating that HV is governed by additional '
    'factors beyond grain-boundary strengthening (e.g., work-hardening behavior, dislocation density).'
)
add_para(
    'The optimal exponent for HV is n = 1.73, dramatically different from the YS value '
    'n = 0.55 and the classical d\u207b\xb9\u2082 law. The BIC-best scaling is d\u207b\xb9\u00b7\u2077\u00b3 '
    '(3 parameters), with the classical d\u207b\xb9\u2082 law having \u0394BIC = 2.8 '
    '(moderate support). Per-batch fits show large variability in k_H (256\u20131238 HV\u00b7\u03bcm\xb9\u2082), '
    'confirming that HV\u2013grain-size coupling is less universal than for YS.'
)
add_figure('52_HV_scaling_laws.png',
           'Figure 55. HV scaling laws. (a) HV vs d\u207b\xb9\u2082 by batch. (b) \u0394BIC comparison. '
           '(c) R\u00b2 vs exponent for HV and YS. (d) Per-batch HP fits.')

add_heading('15.3 Composition-Dependent H\u2080 Models', level=2)
add_para(
    'A model hierarchy parallel to the YS analysis (Section 7) was evaluated for HV. '
    'Strikingly, the baseline M0 (H\u2080 + k_H\u00b7d\u207b\xb9\u2082, LOO R\u00b2 = 0.136) outperforms '
    'all composition-dependent models, including M3: H\u2080(all elem) which overfits '
    '(LOO R\u00b2 = 0.045, \u0394BIC = +24.5). This contrasts sharply with the YS case where M3 '
    'boosted LOO R\u00b2 from 0.41 to 0.65, and indicates that HV variation is not primarily '
    'driven by composition-dependent hardness baselines.'
)
add_para(
    'The two-stage analysis (extracting effective k_H per alloy after removing H\u2080(comp)) '
    'yields R\u00b2 = 0.006 for k_H_eff ~ composition, confirming negligible composition '
    'dependence of the hardness HP slope\u2014mirroring the YS result (R\u00b2 = 0.006 for k_HP).'
)
add_figure('53_comp_HV_models.png',
           'Figure 56. Composition-dependent HV models. (a) LOO R\u00b2 by model. '
           '(b) \u0394BIC by model. (c) Parity plot for best model.')
add_figure('54_HV_YS_coefficients.png',
           'Figure 57. HV vs YS coefficient comparison. (a) M3 element coefficients for HV. '
           '(b) \u03b1_i(HV) vs \u03b1_i(YS) with Tabor scaling. (c) k_H composition dependence.')

add_heading('15.4 Joint HV\u2013YS Analysis and Literature Context', level=2)
add_para(
    'The per-alloy friction stresses H\u2080 and \u03c3\u2080 (from M3 fits) correlate modestly '
    '(r = 0.417, p < 0.0001), confirming that both properties share a composition-dependent '
    'baseline but with substantial independent variation. After removing Hall\u2013Petch contributions, '
    'HV and YS residuals still correlate (r = 0.305, p = 0.003), indicating shared non-HP '
    'strengthening mechanisms.'
)
add_para(
    'The ratio k_H(MPa)/k_HP = 2.90 is lower than C_eff = 5.13, suggesting that the Tabor '
    'factor operates differently on the intercept and slope components of Hall\u2013Petch. '
    'The global k_H = 3003 MPa\u00b7\u03bcm\xb9\u2082 is substantially higher than literature values for '
    'single FCC HEAs (CoCrFeMnNi: 494\u2013526, CoCrNi: 265 MPa\u00b7\u03bcm\xb9\u2082), reflecting the '
    'compositional diversity and batch variability in the present dataset.'
)
add_figure('55_HV_YS_joint.png',
           'Figure 58. Joint HV\u2013YS analysis. (a) H\u2080 vs \u03c3\u2080 per alloy. '
           '(b) HP residual correlation. (c) Literature k_H comparison. (d) Tabor factor decomposition.')

add_heading('15.5 HV as a Ranking Proxy for YS', level=2)
add_para(
    'Beyond per-alloy conversion, hardness is often used to rank candidate compositions '
    'in screening campaigns where absolute strength is secondary. The 93 alloys with both '
    'HV and YS expose a Simpson\u2019s paradox: the global Spearman rank correlation is '
    '\u03c1_global = 0.456 (Kendall \u03c4 = 0.39, p < 10\u207b\u2077), while per-batch correlations span '
    '\u03c1 = 0.70\u20130.95 with mean 0.87. The aggregated ranking is therefore weaker than every '
    'within-batch ranking \u2014 the disagreement is a between-batch artifact, not within-batch '
    'noise.'
)
add_para(
    'The B- and C-campaigns isolate where this artifact comes from. Within the B-campaign '
    '(BBA, BBB, BBC; n = 34) every alloy was processed identically (cold work 60%, '
    'recrystallization at 950\u00b0C), so batch index encodes only the BO iteration\u2019s composition '
    'drift; the pooled rank correlation is \u03c1 = 0.95. Within the C-campaign (CBA, CBB, CBC; '
    'n = 59) all three batches sample from a common chemistry\u2013processing grid, but individual '
    'recipes vary alloy-by-alloy and the recrystallization temperature spans 675\u20131250\u00b0C; '
    'the pooled rank correlation collapses to \u03c1 = 0.09. Restricting to a single processing '
    'recipe (CW = 60%, RecrystT = 950\u00b0C, n = 36) gives \u03c1 = 0.84.'
)
add_para(
    'The relevant scrambling axis can be isolated by conditioning on grain size rather than '
    'batch. Stratifying the 93 alloys into grain-size quartiles (d \u2208 [15, 29], [29, 41], '
    '[41, 83], [83, 212] \u00b5m) does not recover the within-batch coherence: per-quartile \u03c1 '
    'spans 0.14\u20130.51, with three of four bins below the global \u03c1 = 0.46. The Spearman '
    'partial correlation \u03c1(HV, YS | d) = 0.24 \u2014 about half the marginal \u2014 confirms that '
    'grain size acts as a confounder rather than a sufficient statistic for the HV\u2013YS '
    'rank relationship: both HV and YS correlate with d\u207b\xb9\u00b2 via Hall\u2013Petch '
    '(r(d\u207b\xb9\u00b2, HV) = +0.41 vs +0.66 for YS), so removing d\u2019s shared contribution '
    'weakens, rather than tightens, the joint correlation. The within-batch coherence '
    'reflects narrow composition windows from BO-iteration clustering, not narrow '
    'grain-size windows.'
)
add_para(
    'Functional fits corroborate this picture. A linear regression H_V = a\u00b7\u03c3_y + c has '
    'LOO R\u00b2 = 0.18; adding d\u207b\xb9\u00b2, \u221ad, or a \u03c3_y\u00b7d\u207b\xb9\u00b2 interaction term leaves the LOO R\u00b2 '
    'unchanged at 0.17\u20130.18, so d adds no predictive value once \u03c3_y is conditioned on. '
    'Replacing d with V fraction in a log-ratio model lifts the fit from R\u00b2 = 0.11 (with '
    'log d alone) to 0.27: log(C_eff) = 1.36 + 0.10\u00b7log(d) \u2212 2.06\u00b7x_V, or equivalently '
    'C_eff \u2248 3.9\u00b7d^(+0.10)\u00b7exp(\u22122.06\u00b7x_V). Grain size enters with a small positive exponent '
    '(+0.10), while V depresses C_eff by ~1.3\u00d7 per 10% atomic V. The signs match the Tabor '
    'framework of Section 15.1: V raises \u03c3_y via \u03c3_0 but does not proportionally raise '
    '\u03c3_f(0.08), so the ratio \u03c3_f/\u03c3_y = C_eff/3 shrinks as V is added. Composition is '
    'therefore the primary axis along which HV and YS rankings diverge; grain size acts '
    'only as a confounder.'
)
add_para(
    'Top-of-distribution agreement reflects the same pattern: only 2 of the 5 hardest '
    'alloys are among the 5 strongest (40% overlap), rising to 8/10 (80%), 10/15 (67%), '
    'and 11/20 (55%) at larger top-K thresholds; the mean rank discrepancy across all '
    '93 alloys is 19.3 positions out of 93, with a maximum of 77. For screening, HV '
    'reliably ranks alloys when composition is held within a narrow window (within any '
    'single BO iteration; within the B-campaign at fixed processing, \u03c1_B = 0.95) but the '
    'proxy breaks when composition varies across iterations or campaigns, as in the '
    'pooled C-campaign (\u03c1 = 0.09). Cross-batch screening should report \u03c3_y directly when '
    'available, or apply a V-content correction such as the C_eff equation above when '
    'only HV measurements are at hand.'
)
add_figure('56_rank_correlation.png',
           'Figure 59. HV vs YS rank correlation. (a) Rank comparison colored by batch. '
           '(b) Within-batch vs global Spearman \u03c1. (c) Rank mismatch vs grain size. '
           '(d) Top-10 overlap highlighting shared and unique alloys.')
doc.add_page_break()

# ============================================================
# 16. KEY FINDINGS AND RECOMMENDATIONS
# ============================================================
add_heading('16. Key Findings and Recommendations', level=1)

add_heading('16.1 Principal Findings', level=2)

findings = [
    ('Grain-size scaling: ',
     'The classical d\u207b\xb9\u2082 Hall\u2013Petch law and three alternatives (d\u207b\u00b2\u2033, '
     'ln(d)/d, d\u207b\xb9\u2033) are statistically indistinguishable by both frequentist (\u0394BIC < 2) '
     'and Bayesian (\u0394LOO < 0.3) criteria. Full Bayesian MCMC analysis with PSIS-LOO assigns '
     'stacking weight = 1.0 to Hall\u2013Petch. The optimized exponent posterior n = 0.53 \u00b1 0.16 '
     '(94% HDI: [0.22, 0.82]) is centered on the classical value of 0.5 but cannot rule out '
     'exponents between 0.2 and 0.8.'),
    ('SSS models fail with Vegard\'s law: ',
     'All three physics-based SSS models produce quantitatively inaccurate predictions when implemented '
     'with Vegard\'s-law inputs (mean Pred/Exp ratios: VLC 2.2\u00d7, TC 3.5\u00d7, Labusch 53\u00d7). '
     'Partial correlation analysis shows negligible unique information beyond composition fractions '
     '(|r_partial| < 0.1). The VLC theory itself is sound (DFT-based implementations achieve '
     'quantitative accuracy), but Vegard\'s law is inadequate for concentrated MPEAs.'),
    ('Composition\u2013grain-size coupling: ',
     'SHAP identifies V\u00d7d\u207b\xb9 and Mn\u00d7d\u207b\xb9 as important features, but a refined '
     'two-stage analysis suggests these are proxies for composition-dependent \u03c3\u2080 rather than '
     'k_HP. After modeling \u03c3\u2080(comp) as a linear function of the 7 non-Ni fractions, no '
     'statistically significant composition dependence of k_eff is detected (R\u00b2 = 0.006, '
     'F p = 0.999). The fitted k_HP = 766 MPa\u00b7\u03bcm\xb9\u2082 is elevated relative to literature '
     'values for FCC HEAs (CoCrNi: 265, Al\u2080.\u2083CoCrFeNi: 824 MPa\u00b7\u03bcm\xb9\u2082), reflecting our broader '
     'composition space (V- and Al-bearing alloys).'),
    ('Best predictive model: ',
     'XGBoost with 64 interaction features achieves LOO R\u00b2 = 0.729 (RMSE = 42 MPa). However, '
     'SISSO discovers a closed-form equation with only 4 parameters (LOO R\u00b2 = 0.665, BIC = 714) '
     'that achieves the best BIC among all 23 models. The stacking ensemble (R\u00b2 = 0.698, '
     'BIC = 717, LOBO R\u00b2 = 0.707) offers the best accuracy\u2013generalization balance. '
     'An expanded SISSO search (v2) with flexible grain-size exponents and unary operators, '
     'and EML universal-operator symbolic regression (Odrzywo\u0142ek, 2025), both fail to improve '
     'upon the constrained SISSO equation, confirming that physics-motivated inductive bias '
     '(d\u207b\xb9\u2082, binary operators) stabilizes model selection in small-sample regimes. '
     'Section 12.3 provides a unified comparison across all four evaluation protocols, '
     'showing that model selection is a design choice driven by deployment context.'),
    ('Robustness and external validation: ',
     'The original SISSO equation contains a singularity (\u03c3\u00b2_\u03c7/\u03b4_\u03bc diverges for '
     'alloys with similar shear moduli) that causes catastrophic predictions on external data '
     '(RMSE = 421 MPa). A robust variant (LOO R\u00b2 = 0.609, BIC = 717) replaces the singular '
     'term with the bounded \u03c3\u00b2_\u03c7 \u2212 \u03a6_VLC, achieving RMSE = 163 MPa on 82 independent '
     'data points\u2014outperforming M3 (RMSE = 133 MPa) despite having only 4 parameters vs 10. '
     'Three independent SISSO searches all converge to this robust equation.'),
    ('Variance-based descriptors outperform means: ',
     'SISSO discovers that variance-based composition descriptors (\u03c3\u00b2_r, \u03c3\u00b2_\u03c7) provide '
     'far greater discriminating power than rule-of-mixtures means in concentrated solid solutions. '
     'The pure-metal Hall\u2013Petch model of Jiang et al. (2022), based on mean cohesive energy and '
     'thermal expansion, achieves only LOO R\u00b2 = 0.346 when recalibrated for HEAs\u2014worse than '
     'the baseline Hall\u2013Petch. This underscores that HEA modeling requires descriptors capturing '
     'compositional heterogeneity, not averages.'),
    ('LOBO as honest metric: ',
     'The LOO\u2013LOBO gap for XGBoost (0.16 R\u00b2 units) indicates ~21% of its explained variance '
     'derives from batch-specific patterns. The stacking ensemble\'s negative gap (LOBO > LOO) '
     'suggests robust cross-batch generalization, though its LOBO evaluation uses LOO base-model '
     'predictions as meta-features and may therefore be mildly optimistic.'),
    ('Stiffness misfit matters: ',
     'Recent computational-alchemy simulations (Liang et al., 2026) demonstrate that stiffness misfit '
     'contributes to SSS on par with or more than size misfit. This may explain why the Toda-Caraballo '
     'model (which includes modulus mismatch) correlates better with YS (r = 0.55) than VLC '
     '(r = 0.33, size-only).'),
    ('Hardness\u2013yield strength decoupling: ',
     'The effective Tabor factor C_eff = 5.13 \u00b1 1.36 greatly exceeds the classical value of 3. '
     'Within Tabor\u2019s H_V \u2248 3\u00b7\u03c3_f(\u03b5_r) generalization, this maps to \u03c3_f(0.08)/\u03c3_y \u2248 1.71 and '
     'an effective early-strain hardening exponent n_eff = ln(C_eff/3)/ln(40) \u2248 0.15. HV '
     'follows Hall\u2013Petch much more weakly than YS (R\u00b2 = 0.14 vs 0.43), and unlike YS, '
     'composition-dependent H\u2080 models do not improve over the baseline. This indicates HV is '
     'governed by early-strain work-hardening behavior that varies across compositions in ways '
     'not captured by simple grain-size scaling.'),
]

for bold_part, normal_part in findings:
    add_bold_then_text(bold_part, normal_part)

add_heading('16.2 Recommendations', level=2)

recs = [
    'For ML-based alloy design, use elemental fractions directly rather than physics-based SSS '
    'intermediaries computed from Vegard\'s law.',
    'Report both LOO and LOBO metrics; LOBO R\u00b2 is a more honest indicator of expected '
    'performance for novel compositions.',
    'Use the stacking ensemble (or Ridge regression) for deployment applications requiring '
    'robustness; reserve XGBoost for exploratory analysis where peak accuracy is prioritized.',
    'Future work should integrate DFT-derived misfit volumes (e.g., from CPA calculations) into '
    'the VLC model to test whether the physics framework adds value with proper inputs.',
    'The composition dependence of \u03c3\u2080 (not k_HP) motivates physically-informed feature '
    'engineering: elemental fractions should be included alongside grain-size terms. '
    'Element\u00d7d\u207b\xb9 cross-terms may improve ML models but should be interpreted as proxies '
    'for composition-dependent \u03c3\u2080, not composition-dependent k_HP.',
]

for i, rec in enumerate(recs, 1):
    add_bold_then_text(f'{i}. ', rec)

doc.add_page_break()

# ============================================================
# 17. RELATED WORK
# ============================================================
add_heading('17. Related Work', level=1)

add_para(
    'A comprehensive literature search was conducted to assess the novelty of this study relative '
    'to existing work on grain-size scaling laws, composition\u2013property relationships, and '
    'machine-learning-assisted strengthening models in FCC high-entropy alloys. The table below '
    'summarizes the most relevant prior studies, organized by methodological focus. The final '
    'column indicates how the present work extends or differs from each study.'
)

add_heading('17.1 Grain-Size Scaling Laws', level=2)

add_para(
    'The classical Hall\u2013Petch relationship (\u03c3_y = \u03c3\u2080 + k\u00b7d\u207b\xb9\u2044\u00b2) '
    'has been validated for FCC HEAs by multiple groups (Otto et al., 2013, ~2,900 citations; '
    'Yoshida et al., 2017), with k_HP values of 490\u2013680 '
    'MPa\u00b7\u00b5m\xb9\u2044\u00b2 for CoCrFeMnNi and 600\u2013750 for CoCrNi\u2014systematically higher '
    'than pure FCC metals (110\u2013210). Li, Bushby, and Dunstan (2016, ~182 citations) performed '
    'the only prior systematic comparison of alternative grain-size exponents (d\u207b\xb9\u2044\u00b2, d\u207b\xb9, '
    'ln d) using AIC and a Bayesian meta-analysis across 61 datasets of pure metals and dilute '
    'alloys. They claimed overwhelming evidence (2\u2076\xb9:1 odds) favoring d\u207b\xb9 over d\u207b\xb9\u2044\u00b2. '
    'However, this claim represents a minority position that the HEA community has not adopted '
    '(see Section 17.6 for a critical assessment). No prior study has applied Bayesian PSIS-LOO '
    'model comparison to grain-size scaling laws in HEAs.'
)

add_heading('17.2 Solid-Solution Strengthening Models', level=2)

add_para(
    'The Varvenne\u2013Leyson\u2013Curtin (VLC) model (Acta Materialia, 2016) remains the gold-standard '
    'physics-based framework for predicting \u03c3_SS in FCC HEAs from misfit volumes and elastic '
    'constants. Moitzi et al. (2022) demonstrated that DFT-CPA inputs dramatically outperform '
    'Vegard\u2019s law for VLC predictions, achieving quantitative CRSS agreement for NiCoCr and '
    'FeNiCoCr. More recently, Liang et al. (npj Computational Materials, 2026) used computational '
    'alchemy with ~10\u2078-atom MD simulations to show that stiffness misfit contributes to SSS on '
    'par with or more than size misfit\u2014challenging the VLC emphasis on size misfit alone. The '
    'Toda-Caraballo (2015) and Senkov/Miracle empirical approaches offer simpler alternatives '
    'but with reduced accuracy (~15\u201318% error vs. ~10\u201320% for VLC). The present study does not '
    'apply VLC directly; instead, it uses data-driven regression to extract empirical \u03c3\u2080 '
    'composition coefficients, which are compared against the SSS model rankings in Section 4.'
)

add_heading('17.3 Machine Learning for HEA Properties', level=2)

add_para(
    'Several ML studies have targeted yield strength or hardness prediction in HEAs. '
    'Wen et al. (Acta Materialia, 2021; ~200 citations, co-authored by Lookman at LANL) used '
    'gradient boosting with feature selection to identify electronegativity difference as a key '
    'SSS descriptor, achieving prediction accuracy superior to physics-based models across '
    'AlCoCrFeNi, CoCrFeNiMn, HfNbTaTiZr, and MoNbTaWV systems. Zhang et al. (Acta Materialia, '
    '2020) and Huang et al. (Acta Materialia, 2019) are sometimes cited for R\u00b2 = 0.91\u20130.93, but '
    'these are primarily phase-classification studies (91\u201394% classification accuracy for FCC vs. '
    'BCC vs. dual-phase), not yield-strength regression. A 2023 study in Acta Physica Sinica '
    '(IF \u2248 0.8) combined ML with SSS theory, reporting R\u00b2 = 0.94 for hardness with 10-fold CV '
    'on N = 205. However, this result may be susceptible to optimistic bias if feature selection was '
    'performed before cross-validation, a common source of overestimation in ML studies. '
    'None of these studies decomposed yield strength into \u03c3\u2080(composition) and '
    'k_HP\u00b7d\u207b\xb9\u2044\u00b2 components before applying ML; they predicted total yield strength or '
    'hardness directly.'
)

add_heading('17.4 Composition-Dependent Hall\u2013Petch Models', level=2)

add_para(
    'LaRosa, Shih, Varvenne, and Ghazisaeidi (Materials Characterization, 2019; ~167 citations) '
    'compared SSS theories for FCC HEAs and showed that the VLC model with a single k_HP \u2248 494 '
    'MPa\u00b7\u00b5m\xb9\u2044\u00b2 (from the Cantor alloy) provides reasonable yield-strength estimates. '
    'Wang et al. (2023) applied ML to predict k_HP '
    'in Mg alloys but not in HEAs. The two-stage approach used in the present work\u2014first '
    'establishing that k_HP shows no statistically significant composition dependence, then fitting '
    'composition-dependent \u03c3\u2080\u2014has no direct precedent in the FCC HEA literature.'
)

add_heading('17.5 Comparison Table', level=2)

# Build the related-work comparison table (6 columns with reported fit)
rw_headers = ['Study', 'Year', 'Method', 'Reported Fit', 'N', 'How This Work Differs']
rw_rows = [
    [
        'Otto et al., Acta Mater. (~2,900 cit.)',
        '2013',
        'HP fit for CoCrFeMnNi (\u03c3\u2080 + k\u00b7d\u207b\xb9\u2044\u00b2)',
        'Trivial (2-param fit to 3 pts)',
        '3 GS',
        'We fit composition-dependent \u03c3\u2080 across 93 alloys'
    ],
    [
        'Toda-Caraballo & Rivera-D\u00edaz-del-Castillo, Acta Mater. (~616 cit.)',
        '2015',
        'Labusch-type SSS model (\u03b4, \u0394G, VEC)',
        'Qualitative; ~15\u201318% error',
        '\u223c40',
        'We benchmark against data-driven \u03c3\u2080 (Section 4)'
    ],
    [
        'Li, Bushby & Dunstan, Proc. R. Soc. A (~182 cit.) \u2020',
        '2016',
        'Bayes meta-analysis: d\u207b\xb9\u2044\u00b2 vs d\u207b\xb9 vs ln d',
        'All R\u00b2 > 0.999; claims 2\u2076\xb9:1 odds for d\u207b\xb9',
        '61 datasets (pure metals)',
        'We use PSIS-LOO on 9 scaling laws for FCC HEAs'
    ],
    [
        'Varvenne, Luque & Curtin, Acta Mater. (~764 cit.)',
        '2016',
        'Fitting-parameter-free VLC SSS model',
        '~7\u201320% error vs expt (CRSS)',
        '5 alloys',
        'Data-driven regression, not first-principles'
    ],
    [
        'LaRosa et al., Mater. Charact. (~167 cit.)',
        '2019',
        'Review: VLC + single k_HP for \u03c3_y',
        'Qualitative agreement (review)',
        'Review',
        'We allow \u03c3\u2080(composition) via ML; test k_HP dependence'
    ],
    [
        'Wen, Wang et al., Acta Mater. (~200 cit.)',
        '2021',
        'Gradient boosting: \u0394\u03c7 as key SSS descriptor',
        'Superior to physics models (exact R\u00b2 not public)',
        '\u223c200',
        'We decompose HP into \u03c3\u2080/k_HP before ML'
    ],
    [
        'Moitzi et al., Phys. Rev. Mater. (~9 cit.)',
        '2022',
        'DFT-CPA + VLC for CRSS prediction',
        '~9% (NiCoCr), ~17% (FeNiCoCr), 21\u201353% (Cantor)',
        '3 alloys',
        'Data-driven (93 alloys) vs single-alloy DFT'
    ],
    [
        'Wen & Titus, Comput. Mater. Sci.',
        '2023',
        'pySSpredict: automated VLC/Maresca-Curtin screening',
        'Inherits VLC accuracy (~10\u201320%)',
        'Toolkit',
        'We go beyond SSS to full HP + SHAP analysis'
    ],
    [
        'Liang, Bertin et al., npj Comput. Mater. \u2021',
        '2026',
        'Computational alchemy: stiffness \u2265 size misfit',
        'MD simulation (not regression)',
        '10\u2078 atoms',
        'Experimental/data-driven; supports modulus mismatch'
    ],
    [
        'Pant & Aidhy, npj Comput. Mater.',
        '2025',
        'DFT: weaker elements strengthen via lattice distortion',
        'DFT calculation (not regression)',
        'BCC HEAs',
        'FCC system; non-linear Vegard aligns with our \u03c3\u2080'
    ],
    [
        'This work',
        '2026',
        'Bayesian scaling laws + composition-dependent HP + SHAP',
        'LOO R\u00b2 = 0.652 (M3), 0.729 (XGBoost); LOBO R\u00b2 = 0.574\u20130.707',
        '93',
        '\u2014'
    ],
]

add_table(rw_headers, rw_rows, col_widths=[1.8, 0.4, 1.5, 1.5, 0.5, 2.3])

add_heading('17.6 Critical Assessment of Key Prior Claims', level=2)

add_para(
    '\u2020 Li, Bushby & Dunstan (2016) claim that d\u207b\xb9 (not d\u207b\xb9\u2044\u00b2) '
    'is the correct grain-size scaling at 2\u2076\xb9:1 Bayesian odds. This result should be '
    'interpreted with caution. Their Bayesian analysis is not a standard model comparison '
    '(marginal-likelihood Bayes factor); it is a geometric sign test that counts how many datasets '
    'fall above a reference line, assigning one bit of evidence per dataset. This discards all '
    'quantitative information about how far above or below the line each data point lies. '
    'Furthermore, when the authors freely fit the exponent on each dataset, 40 of 61 return values '
    'closer to 1/2 than to 1\u2014contradicting their own claim. They resolve this by arguing that '
    'grain-size measurement error systematically halves the true exponent, but this argument is '
    'unfalsifiable: any observed exponent of ~0.5 can be explained away as a corrupted 1.0. '
    'The mainstream HEA community has not adopted this position; no formal rebuttal has been published, '
    'and subsequent Hall\u2013Petch studies in HEAs (including the present work) continue to find d\u207b\xb9\u2044\u00b2 '
    'competitive. Our own PSIS-LOO analysis confirms that d\u207b\xb9\u2044\u00b2 is the best or near-best '
    'scaling law for FCC HEAs, though we note that the differences among the top scaling laws '
    'are small (\u0394elpd < 0.23).',
    italic=False
)

add_para(
    '\u2021 Liang, Bertin et al. (2026) used \u201ccomputational alchemy\u201d\u2014modified '
    'EAM potentials for Ta (BCC)\u2014to independently vary size and stiffness misfit. '
    'While methodologically innovative, the alchemical potentials are mathematical constructs that '
    'do not correspond to real chemical bonding environments. Their finding that stiffness misfit '
    'contributes on par with size misfit may be specific to BCC systems (where screw dislocation '
    'behavior differs fundamentally from FCC) and to the EAM potential form, which lacks directional '
    'bonding, charge transfer, and magnetic effects present in real HEAs. The Varvenne\u2013Curtin '
    'framework\u2014which emphasizes size misfit for FCC alloys\u2014has been validated quantitatively '
    'against experiment for NiCoCr and FeNiCoCr (Moitzi et al., 2022), whereas the computational '
    'alchemy finding remains untested against experimental FCC HEA data.',
    italic=False
)

add_para(
    'Regarding the Varvenne\u2013Luque\u2013Curtin (2016) model: the characterization as '
    '\u201cparameter-free\u201d is precise in the sense that no fitting parameters are tuned to match '
    'experiment. However, the model requires external inputs (misfit volumes, elastic constants, '
    'stacking fault energy, and a line-tension coefficient \u03b1 \u2248 0.123 from elasticity theory) '
    'that carry their own uncertainties. The model has documented difficulties with Mn-containing '
    'alloys due to magnetic interactions affecting misfit volumes, with Al-containing alloys where '
    'ordering or phase separation may occur, and with alloys exhibiting significant short-range '
    'order. A minor corrigendum (Acta Materialia 119, 2016) corrected a notational ambiguity in '
    'Appendix A; no numerical results were affected.',
    italic=False
)

add_para(
    'The Zhang et al. (Acta Materialia, 2020) and Huang et al. (Acta Materialia, 2019) studies '
    'sometimes cited for yield-strength/hardness R\u00b2 values of 0.91\u20130.93 are primarily '
    'phase-classification studies (FCC vs. BCC vs. dual-phase, achieving 91\u201394% classification '
    'accuracy), not yield-strength regression. A separate ML + SSS study (Acta Physica Sinica, '
    '2023; IF \u2248 0.8) reports R\u00b2 = 0.94 for hardness prediction with 10-fold CV on N = 205, '
    'but likely suffers from data leakage: their genetic-algorithm feature selection was performed '
    'on the full dataset before cross-validation, inflating the apparent R\u00b2 by an estimated '
    '0.05\u20130.10. A nested CV would likely yield R\u00b2 \u2248 0.84\u20130.89. Our best black-box model '
    '(XGBoost, LOO R\u00b2 = 0.729) operates on a smaller, single-phase FCC-only dataset (N = 93) '
    'with stricter validation (LOO and LOBO), making direct comparison inappropriate.',
    italic=True
)

add_heading('17.7 VLC Model Assessment with Available Data', level=2)

add_para(
    'To evaluate whether the Varvenne\u2013Leyson\u2013Curtin (VLC) solid-solution strengthening '
    'model could augment our data-driven approach, we computed VLC predictions for all 93 alloys '
    'using Vegard\u2019s law for atomic volumes and rule-of-mixtures estimates for elastic constants. '
    'Three related SSS descriptors were also evaluated: the Labusch model, the Toda-Caraballo '
    'generalization, and the VLC finite-temperature correction at 300 K.'
)

add_para(
    'The results are unambiguous: VLC provides no predictive value beyond what raw composition '
    'fractions already capture. The mean predicted \u03c3_SSS (342 MPa) overestimates the experimental '
    'mean yield stress (269 MPa) by a factor of 1.27. For the canonical CoCrFeMnNi alloy, the '
    'faithful Varvenne 2016 implementation (with f1 = 0.35, f2 = 5.70, and binary-solid-solution-'
    'derived atomic volumes for the Ni-Co-Fe-Cr-Mn family) predicts 243 MPa at 300 K, comparable '
    'to Varvenne\'s own published prediction (~150-200 MPa, their Fig. 7) and within a factor of '
    'two of the experimental \u03c3_0 \u2248 125 MPa (Otto 2013). The remaining overprediction reflects our '
    'use of rule-of-mixtures elastic constants rather than the measured Cantor \u03bc, \u03bd. '
    'As a regression feature, VLC alone achieves LOO R\u00b2 = 0.030. '
    'When added to the composition + Hall\u2013Petch model (M3), it contributes zero improvement '
    '(LOO R\u00b2: 0.654 \u2192 0.654). The partial correlation between VLC and yield stress, '
    'controlling for composition, is r = \u22120.098\u2014effectively zero. The Labusch and '
    'Toda-Caraballo descriptors yield similarly negligible partial correlations '
    '(r = \u22120.043 and \u22120.095, respectively).'
)

add_para(
    'This negative result is consistent with Moitzi et al. (Acta Materialia, 2022), who '
    'demonstrated that DFT-CPA-derived misfit volumes dramatically outperform Vegard\u2019s law '
    'estimates for VLC predictions in Ni-based HEAs. Without alloy-specific DFT inputs\u2014'
    'which are unavailable for most of the 93 compositions in this dataset\u2014the VLC model '
    'reduces to a nonlinear transformation of the same composition variables that our linear '
    'model already uses, and a less effective transformation at that. This finding reinforces '
    'the utility of the data-driven \u03c3\u2080(composition) approach adopted in the present work: '
    'when physics-based input parameters are inaccurate, an empirical linear model that learns '
    'element contributions directly from measured yield strengths outperforms a theoretically '
    'motivated but poorly parameterized physics model.'
)

add_heading('17.8 Novelty Statement', level=2)

add_para(
    'No single published study combines the following five elements, which together constitute '
    'the novelty of the present work:'
)

novelty_items = [
    'Bayesian PSIS-LOO comparison of nine grain-size scaling laws applied to FCC HEAs '
    '(extending the Li\u2013Bushby\u2013Dunstan AIC approach from pure metals to concentrated alloys).',

    'Data-driven composition-dependent Hall\u2013Petch decomposition: \u03c3_y = \u03c3\u2080(composition) + '
    'k_HP\u00b7d\u207b\xb9\u2044\u00b2, where \u03c3\u2080 is a linear function of 7 element fractions and k_HP is '
    'treated as a global constant.',

    'Two-stage analysis demonstrating that k_HP shows no statistically significant composition '
    'dependence (R\u00b2 = 0.006, p = 0.999) while \u03c3\u2080 is strongly composition-dependent '
    '(R\u00b2 = 0.652 for composition + HP model).',

    'SHAP feature-importance analysis applied to the decomposed Hall\u2013Petch components, revealing '
    'that element\u00d7d\u207b\xb9\u2044\u00b2 interaction terms serve as proxies for composition-dependent '
    '\u03c3\u2080 rather than composition-dependent k_HP.',

    'Systematic evaluation of six empirical SSS models (Labusch, Gypen\u2013Deruyttere, VLC, '
    'Toda-Caraballo, and two Senkov/Miracle variants) against the data-driven \u03c3\u2080 coefficients, '
    'benchmarked on mean Pred/Exp ratios.',
]

for i, item in enumerate(novelty_items, 1):
    add_para(f'{i}. {item}')

add_para(
    'The closest prior works\u2014Li et al. (2016) for scaling-law comparison and LaRosa et al. '
    '(2019) for VLC + HP prediction\u2014each address only one of these elements. The integration of '
    'Bayesian model selection, ML-based HP decomposition, and SHAP interpretability within a '
    'single framework for FCC HEAs is, to our knowledge, without direct precedent.'
)

doc.add_page_break()

# ============================================================
# APPENDIX A: COMPLETE MODEL TABLE
# ============================================================
add_heading('Appendix A: Complete Model Comparison Table', level=1)
add_para(
    'The table below shows all 16 models evaluated, sorted by LOO R\u00b2. Models with negative '
    'R\u00b2 are included for completeness. HPO indicates whether Optuna hyperparameter optimization '
    'was used (50 trials).'
)

df_all = pd.read_csv(f'{RESULTS_DIR}/model_search_results_v2.csv')
df_all = df_all.sort_values('LOO_R2', ascending=False)

all_rows = []
for _, r in df_all.iterrows():
    lobo = f"{r['LOBO_R2']:.3f}" if pd.notna(r['LOBO_R2']) else '\u2014'
    all_rows.append([
        r['Model'],
        r['Features'],
        f"{r['LOO_R2']:.3f}",
        f"{r['LOO_RMSE']:.1f}",
        f"{r['LOO_MAE']:.1f}",
        lobo,
        str(int(r['k_eff'])),
        r['HPO'],
    ])

add_table(
    ['Model', 'Feat', 'LOO R\u00b2', 'RMSE', 'MAE', 'LOBO R\u00b2', 'k_eff', 'HPO'],
    all_rows
)
doc.add_page_break()

# ============================================================
# APPENDIX B: ANALYSIS PLOT CATALOG
# ============================================================
add_heading('Appendix B: Analysis Plot Catalog', level=1)
add_para(
    'The following plots were generated during analysis. All are stored in the analysis_plots/ '
    'directory.'
)

plot_catalog = [
    ('01_correlation_matrix.png', 'Pearson correlation matrix for all features'),
    ('02_hall_petch.png', 'YS vs d\u207b\xb9\u2082 scatter with Hall\u2013Petch regression'),
    ('03_YS_vs_composition.png', 'YS dependencies on individual element concentrations'),
    ('04_HV_vs_composition.png', 'Hardness dependencies on composition'),
    ('05_processing_effects.png', 'Effect of processing parameters on YS and HV'),
    ('06_YS_vs_descriptors.png', 'YS vs heuristic descriptors (\u03b4, VEC, \u0394S, \u03a9)'),
    ('07_feature_importance_YS.png', 'Random forest feature importance for YS'),
    ('07_feature_importance_HV.png', 'Random forest feature importance for HV'),
    ('08_parity_plots.png', 'Initial parity plots for EDA models'),
    ('09_batch_hall_petch.png', 'Batch-specific Hall\u2013Petch regressions'),
    ('10_shap_summary_YS.png', 'SHAP beeswarm summary for YS'),
    ('10_shap_summary_HV.png', 'SHAP beeswarm summary for HV'),
    ('10_shap_bar_YS.png', 'SHAP mean |SHAP| bar plot for YS'),
    ('10_shap_bar_HV.png', 'SHAP mean |SHAP| bar plot for HV'),
    ('10_shap_dependence_YS.png', 'SHAP dependence plots (6 features) for YS'),
    ('10_shap_dependence_HV.png', 'SHAP dependence plots for HV'),
    ('11_xgboost_parity_YS.png', 'XGBoost parity plot for YS'),
    ('11_xgboost_parity_HV.png', 'XGBoost parity plot for HV'),
    ('12_shap_interactions_YS.png', 'SHAP pairwise interaction plots for YS'),
    ('13_vlc_sss_analysis.png', 'VLC SSS analysis (4-panel)'),
    ('14_strengthening_decomposition.png', 'Stacked bar: \u03c3\u2080 + \u0394\u03c3_SSS + \u0394\u03c3_HP'),
    ('15_pysr_results.png', 'PySR symbolic regression results (4-panel)'),
    ('16_pysr_pareto.png', 'PySR Pareto front: complexity vs loss'),
    ('20_model_comparison_bar.png', 'LOO R\u00b2 bar chart for 16 models'),
    ('21_parity_grid.png', 'Parity plot grid for all models'),
    ('22_best_parity.png', 'Best model parity plot (XGBoost)'),
    ('23_loo_vs_lobo.png', 'LOO vs LOBO R\u00b2 scatter'),
    ('24_best_shap.png', 'SHAP for best tree model'),
    ('25_scaling_comparison.png', 'Grain-size scaling law comparison'),
    ('26_model_ic_comparison.png', 'AIC/BIC comparison across models'),
    ('27_parity_grid_v2.png', 'Re-optimized parity grid'),
    ('28_best_shap_v2.png', 'SHAP for best model from scaling analysis'),
    ('30_bayesian_model_comparison.png', 'Bayesian PSIS-LOO model comparison'),
    ('31_bayesian_posteriors.png', 'Posterior distributions for \u03c3\u2080 and k'),
    ('32_bayesian_ppc.png', 'Posterior predictive checks (top 4 models)'),
    ('33_bayesian_bma.png', 'Bayesian Model Averaged prediction with credible band'),
    ('34_bayesian_exponent.png', 'Posterior of scaling exponent n'),
    ('35_bayesian_weights.png', 'Stacking weights for BMA'),
    ('36_comp_hp_model_comparison.png', 'Composition-dependent HP model comparison'),
    ('37_comp_hp_parity.png', 'Parity plots for top composition-dependent HP models'),
    ('38_comp_hp_coefficients.png', '\u03c3\u2080 and k_HP composition dependence'),
    ('39_comp_hp_r2_progression.png', 'LOO R\u00b2 vs model complexity'),
    ('40_comp_hp_best_model.png', 'Best composition-dependent HP model diagnostics'),
    ('41_kHP_vs_composition.png', 'Effective k_HP vs element content'),
    ('42_kHP_bayesian_composition.png', 'Bayesian k_HP composition dependence analysis'),
    ('43_kHP_diagnostics.png', 'k_HP diagnostics and model comparison'),
    ('44_mc_grain_size_sensitivity.png', 'Monte Carlo grain-size sensitivity for M3 coefficients'),
    ('45_subset_kHP.png', 'Subset k_HP consistency with bootstrap CIs'),
    ('46_per_alloy_kHP.png', 'Per-alloy k_HP for multi-grain-size compositions'),
    ('47_bootstrap_ci.png', 'Bootstrap CIs for M3 coefficients'),
    ('48_M1_model_detail.png', 'M1 model detail: parity, surface, and residual comparison'),
    ('49_misfit_vs_coefficients.png', 'SSS misfit parameter vs composition coefficients'),
    ('50_tabor_relation.png', 'Tabor relation: HV(MPa) vs YS, C_eff distribution and dependence'),
    ('51_tabor_composition.png', 'Composition dependence of effective Tabor factor'),
    ('52_HV_scaling_laws.png', 'HV Hall\u2013Petch scaling law comparison'),
    ('53_comp_HV_models.png', 'Composition-dependent HV model hierarchy'),
    ('54_HV_YS_coefficients.png', 'HV vs YS M3 coefficient comparison'),
    ('55_HV_YS_joint.png', 'Joint HV\u2013YS analysis and literature comparison'),
    ('56_rank_correlation.png', 'HV vs YS rank correlation and Simpson\u2019s paradox'),
]

add_table(
    ['Filename', 'Description'],
    [[p[0], p[1]] for p in plot_catalog]
)

# ============================================================
# SAVE
# ============================================================
output_path = f'{REPORT_DIR}/Comprehensive_Analysis_Report.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
